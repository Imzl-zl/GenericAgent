#!/usr/bin/env python3
"""docx_utils.py — 文档转换工具脚本（GA L3 原生 .py 形态，2026-08-15）。

架构定位（对齐社区 pandoc skill 案例 + GA SOP 规范）：
    - 指令层 document_conversion_sop.md 教"怎么选/怎么组合"，本脚本封装
      "高复用、逻辑复杂、不想每次重新推理"的处理流程（GA L3 .py 定义）。
    - pandoc 优先：md/html→docx、docx→md、pptx/epub 全走 pandoc；
      pandoc 不支持的（PDF 生成、.doc/.xls 老格式）由 SOP 指引 LibreOffice，
      本脚本不重复实现。
    - --help 即配方：每个子命令自解释参数，模型跑 help 即得用法。

子命令：
    make-template   按规格/预设生成 pandoc --reference-doc 模板(样式定制核心)
    md-to-docx      md/html/txt → docx（模板 + toc/编号/高亮/GFM 全参数）
    verify          转换后重读验证（段落/表格/字体，闭环不靠模型自觉）

用法示例：
    python docx_utils.py make-template -o /tmp/ref.docx --preset cn
    python docx_utils.py make-template -o /tmp/ref.docx --spec '{"heading1": {"font": "黑体", "size": "四号"}}'
    python docx_utils.py md-to-docx 输入.md outputs/输出.docx --toc
    python docx_utils.py verify outputs/输出.docx --expect-tables 2
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any

# ---------------------------------------------------------------------------
# 常量与样式规格解析
# ---------------------------------------------------------------------------

# 中文常用字号 → pt（确定性映射，支持"四号"这类叫法）。
CN_FONT_SIZE_PT = {
    "初号": 42, "小初": 36, "一号": 26, "小一": 24, "二号": 22, "小二": 18,
    "三号": 16, "小三": 15, "四号": 14, "小四": 12, "五号": 10.5, "小五": 9,
    "六号": 7.5, "小六": 6.5, "七号": 5.5, "八号": 5,
}

# 用户语义键 → pandoc reference.docx styleId 集合(2026-08-15 审查 A-2):
# pandoc 正文段落实际使用 FirstParagraph(首段)/BodyText(后续段), Normal
# 几乎不被段落引用(实测: 三段正文 = FirstParagraph/BodyText/BodyText)。
# 语义键必须展开为实际使用集合, 否则用户"正文宋体小四"只改到 Normal,
# 产物正文纹丝不动。直写 styleId 则保持单点精确控制。
SEMANTIC_STYLE_MAP = {
    "title": ("Title",),
    "heading1": ("Heading1",), "heading2": ("Heading2",), "heading3": ("Heading3",),
    "heading4": ("Heading4",), "heading5": ("Heading5",), "heading6": ("Heading6",),
    "body": ("Normal", "BodyText", "FirstParagraph"),
    "bodytext": ("BodyText", "FirstParagraph"),
    "firstparagraph": ("FirstParagraph",),
    "table": ("Table",), "compact": ("Compact",), "blocktext": ("BlockText",),
    "sourcecode": ("SourceCode",),
}

# 语义键覆盖的 styleId 全集: 缺失时 make-template 自动补建(pandoc 默认
# 模板无 SourceCode; 语义键=意图, 意图必须落地)。直写 styleId 不在本集合
# 时缺失仅告警——拼写错误显式暴露, 不静默造样式。
AUTO_ENSURE_STYLES = {sid for group in SEMANTIC_STYLE_MAP.values() for sid in group}


# markdown 结构启发式(与 worker legacy_instrument._looks_like_markdown 保持
# 一致, 2026-08-15 审查 A-1: txt 内容判定用, 决定 pandoc 临时源扩展名)。
_MARKDOWN_HINTS = (
    re.compile(r"^#{1,6}\s", re.M),
    re.compile(r"^\s*\|.*\|\s*$", re.M),
    re.compile(r"^\s*[-*]\s+\S", re.M),
    re.compile(r"\*\*[^*]+\*\*"),
    re.compile(r"^```", re.M),
)


def _looks_like_markdown(content: str) -> bool:
    """启发式判定文本是否 markdown 结构(标题/表格/列表/粗体/代码块)。"""
    if not content or not content.strip():
        return False
    head = content[:4000]
    return any(p.search(head) for p in _MARKDOWN_HINTS)

# 中文排版预设（社区 pandoc_docx_template 标准，实测验证）：
# 正文宋体小四+首行缩进2字符+1.5倍行距；H1 黑体小二加粗段前段后24磅；
# H2 黑体三号；H3 黑体小三；H4-6 黑体小四；标题黑体二号居中；
# 表格单元格宋体 10pt；西文一律 Times New Roman。
PRESET_CN = {
    "Normal": {"font": "宋体", "latin": "Times New Roman", "size": 12, "line": 1.5},
    "BodyText": {"font": "宋体", "latin": "Times New Roman", "size": 12, "indent_chars": 2, "line": 1.5},
    "FirstParagraph": {"font": "宋体", "latin": "Times New Roman", "size": 12, "indent_chars": 2, "line": 1.5},
    "Compact": {"font": "宋体", "latin": "Times New Roman", "size": 10},
    "Title": {"font": "黑体", "latin": "Times New Roman", "size": 22, "center": True},
    "Heading1": {"font": "黑体", "latin": "Times New Roman", "size": 18, "bold": True, "before": 24, "after": 24},
    "Heading2": {"font": "黑体", "latin": "Times New Roman", "size": 16, "bold": True, "before": 12, "after": 6},
    "Heading3": {"font": "黑体", "latin": "Times New Roman", "size": 15, "bold": True, "before": 12, "after": 6},
    "Heading4": {"font": "黑体", "latin": "Times New Roman", "size": 12, "bold": True},
    "Heading5": {"font": "黑体", "latin": "Times New Roman", "size": 12, "bold": True},
    "Heading6": {"font": "黑体", "latin": "Times New Roman", "size": 12, "bold": True},
    "BlockText": {"font": "宋体", "latin": "Times New Roman", "size": 12, "line": 1.5},
    "SourceCode": {"font": "Consolas", "latin": "Consolas", "size": 10},
    # 2026-08-15 模板化升级: 表格样式(社区模板标配, 官方 reference-doc 机制
    # 的一部分)——全框线 + 首行表头浅蓝底纹加粗, 由 _apply_table_style 特判
    # 应用(非段落样式: tblPr/tblStylePr)。header_fill 可被 --spec 覆盖。
    "Table": {"font": "宋体", "latin": "Times New Roman", "size": 10,
               "header_fill": "DEEAF6"},
}


def _resolve_pt(size: Any) -> float | None:
    """字号解析: 数字 = pt; 字符串数字 = pt; 中文叫法查表。非法返回 None。"""
    if isinstance(size, (int, float)):
        return float(size) if size > 0 else None
    if isinstance(size, str):
        v = size.strip()
        if v in CN_FONT_SIZE_PT:
            return CN_FONT_SIZE_PT[v]
        try:
            f = float(v)
            return f if f > 0 else None
        except ValueError:
            return None
    return None


def _normalize_spec(spec: dict[str, Any]) -> dict[str, dict[str, Any]]:
    """规格归一化: 语义键/直写 styleId 都映射为 {styleId: {font, latin, size,
    bold, before, after, line, indent_chars, center, header_fill}};
    非法项忽略(宽松)。header_fill 仅 Table 样式使用(表头底纹色, 审查 M1:
    此前不在白名单导致 --spec 静默丢弃, 与 PRESET_CN 注释/SOP 宣传矛盾)。"""
    out: dict[str, dict[str, Any]] = {}
    for key, raw in spec.items():
        if not isinstance(raw, dict):
            continue
        # 语义键展开为实际样式集合; 直写 styleId 单点。
        for style_id in SEMANTIC_STYLE_MAP.get(key.lower(), (key,)):
            clean: dict[str, Any] = {}
            if raw.get("font"):
                clean["font"] = str(raw["font"]).strip()
            if raw.get("latin"):
                clean["latin"] = str(raw["latin"]).strip()
            pt = _resolve_pt(raw.get("size"))
            if pt is not None:
                clean["size"] = pt
            if raw.get("bold") is not None:
                clean["bold"] = bool(raw["bold"])
            for k in ("before", "after"):
                v = raw.get(k)
                if isinstance(v, (int, float)) and v >= 0:
                    clean[k] = float(v)
            v = raw.get("line")
            if isinstance(v, (int, float)) and v > 0:
                clean["line"] = float(v)
            v = raw.get("indent_chars")
            if isinstance(v, (int, float)) and v > 0:
                clean["indent_chars"] = int(v)
            if raw.get("center"):
                clean["center"] = True
            if raw.get("header_fill"):
                clean["header_fill"] = str(raw["header_fill"]).strip()
            if clean:
                out[style_id] = clean
    return out


# ---------------------------------------------------------------------------
# make-template：按规格生成 pandoc reference.docx 模板
# ---------------------------------------------------------------------------

def _find_style(styles_el: Any, style_id: str):
    from docx.oxml.ns import qn
    for st in styles_el.iter(qn("w:style")):
        if (st.get(qn("w:styleId")) or "").lower() == style_id.lower():
            return st
    return None


def _ensure_style(doc: Any, style_id: str, base_id: str = "Normal") -> bool:
    """为缺失的样式补建最小骨架(基于 Normal)。pandoc 默认 reference.docx 无
    SourceCode(实测 12/13), 语义键必须落地——补建后 _apply_style 再应用。"""
    from docx.oxml.ns import qn

    if _find_style(doc.styles.element, style_id) is not None:
        return True
    st = doc.styles.element.makeelement(qn("w:style"), {})
    st.set(qn("w:type"), "paragraph")
    st.set(qn("w:styleId"), style_id)
    for tag, val in (("w:name", style_id), ("w:basedOn", base_id)):
        el = st.makeelement(qn(tag), {})
        el.set(qn("w:val"), val)
        st.append(el)
    # CT_Style 顺序: ... unhideWhenUsed, qFormat(审查: 旧代码 qFormat 在前,
    # 违反 schema 序——strict 校验器会拒, 现按序 append)。
    unhide = st.makeelement(qn("w:unhideWhenUsed"), {})
    st.append(unhide)
    qf = st.makeelement(qn("w:qFormat"), {})
    st.append(qf)
    doc.styles.element.append(st)
    return True


def _apply_table_style(doc: Any, style_id: str, spec: dict[str, Any]) -> bool:
    """Table 样式(表格样式, 非段落): 全框线 0.5pt 灰 + 首行表头底纹加粗。

    社区模板标配(参考 Achuan-2/pandoc_docx_template 等): pandoc 默认模板的
    Table 样式只有单元格内边距, 表格观感=朴素网格; 本函数补 tblBorders(全
    框线)与 tblStylePr firstRow(表头底纹/加粗)。样式缺失时补建 type=table
    (pandoc 默认模板已含 Table, 补建仅为防御)。子元素一律按 OOXML schema
    顺序插入(CT_TblPr: tblBorders 必须在 tblCellMar 之前; CT_TblStylePr:
    tcPr 在 rPr 之前)。"""
    from docx.oxml.ns import qn

    target = _find_style(doc.styles.element, style_id)
    if target is None:
        target = doc.styles.element.makeelement(qn("w:style"), {})
        target.set(qn("w:type"), "table")
        target.set(qn("w:styleId"), style_id)
        name = target.makeelement(qn("w:name"), {})
        name.set(qn("w:val"), style_id)
        target.append(name)
        qf = target.makeelement(qn("w:qFormat"), {})
        target.append(qf)
        doc.styles.element.append(target)
    # 字体(与段落样式同构: 中文=宋体, 西文=Times New Roman)
    latin = spec.get("latin", spec.get("font", "Times New Roman"))
    rPr = target.find(qn("w:rPr"))
    if rPr is None:
        rPr = target.makeelement(qn("w:rPr"), {})
        _insert_style_child_in_order(target, rPr, "rPr")
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)
    rFonts.set(qn("w:eastAsia"), spec.get("font", latin))
    rFonts.set(qn("w:cs"), latin)
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rFonts.get(qn(attr)) is not None:
            del rFonts.attrib[qn(attr)]
    if "size" in spec:
        half = str(int(round(spec["size"] * 2)))
        for tag in ("w:sz", "w:szCs"):
            sz = rPr.find(qn(tag))
            if sz is None:
                sz = rPr.makeelement(qn(tag), {})
                rPr.append(sz)
            sz.set(qn("w:val"), half)
    # 全框线: tblPr/tblBorders(单线 0.5pt 灰)。CT_TblPr 顺序: tblBorders
    # 在 shd/tblLayout/tblCellMar 之前——pandoc 默认 Table 样式已有
    # tblCellMar, 直接 append 会违反 schema 序, 必须插到它前面。
    tblPr = target.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = target.makeelement(qn("w:tblPr"), {})
        _insert_style_child_in_order(target, tblPr, "tblPr")
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = tblPr.makeelement(qn("w:tblBorders"), {})
        before = tblPr.find(qn("w:tblCellMar"))
        if before is not None:
            before.addprevious(tblBorders)
        else:
            tblPr.append(tblBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tblBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = tblBorders.makeelement(qn(f"w:{edge}"), {})
            tblBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "7F7F7F")
    # 表头行: 底纹 + 加粗。CT_TblStylePr 顺序: tcPr 在 rPr 之前。
    tblStylePr = target.find(qn("w:tblStylePr"))
    if tblStylePr is None:
        tblStylePr = target.makeelement(qn("w:tblStylePr"), {})
        _insert_style_child_in_order(target, tblStylePr, "tblStylePr")
    tblStylePr.set(qn("w:type"), "firstRow")
    tcPr = tblStylePr.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = tblStylePr.makeelement(qn("w:tcPr"), {})
        tblStylePr.insert(0, tcPr)
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = tcPr.makeelement(qn("w:shd"), {})
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), spec.get("header_fill", "DEEAF6"))
    hrPr = tblStylePr.find(qn("w:rPr"))
    if hrPr is None:
        hrPr = tblStylePr.makeelement(qn("w:rPr"), {})
        tcPr.addnext(hrPr)  # tcPr 之后, CT_TblStylePr 顺序
    b = hrPr.find(qn("w:b"))
    if b is None:
        b = hrPr.makeelement(qn("w:b"), {})
        hrPr.append(b)
    b.set(qn("w:val"), "1")
    return True


# OOXML CT_Style 子元素顺序表(insert 定位用)。
# 顺序: name, aliases, basedOn, next, link, autoRedefine, hidden, uiPriority,
# semiHidden, unhideWhenUsed, qFormat, locked, personal, personalCompose,
# personalReply, rsid, pPr, rPr, tblPr, trPr, tcPr, tblStylePr。
_STYLE_CHILD_RANK = {
    "name": 0, "aliases": 1, "basedOn": 2, "next": 3, "link": 4,
    "autoRedefine": 5, "hidden": 6, "uiPriority": 7, "semiHidden": 8,
    "unhideWhenUsed": 9, "qFormat": 10, "locked": 11, "personal": 12,
    "personalCompose": 13, "personalReply": 14, "rsid": 15,
    "pPr": 16, "rPr": 17, "tblPr": 18, "trPr": 19, "tcPr": 20,
    "tblStylePr": 21,
}


def _insert_style_child_in_order(target: Any, el: Any, tag: str) -> None:
    """按 CT_Style 子元素顺序把 el 插入 target(样式元素)。

    旧代码 insert(0) 会把 pPr/rPr 插到 w:name 之前, 违反 schema 序(strict
    校验器/LibreOffice 可能拒); 直接 append 又会把 tblBorders 放到
    tblCellMar 之后。本函数按秩表找第一个秩更大的已有子元素, 插到它前面;
    没有则 append。"""
    from docx.oxml.ns import qn

    rank = _STYLE_CHILD_RANK.get(tag)
    if rank is None:
        target.append(el)
        return
    for child in target:
        local = child.tag.rsplit("}", 1)[-1]
        child_rank = _STYLE_CHILD_RANK.get(local)
        if child_rank is not None and child_rank > rank:
            child.addprevious(el)
            return
    target.append(el)


def _apply_style(doc: Any, style_id: str, spec: dict[str, Any]) -> bool:
    """对 styles.xml 中指定 styleId 应用字体/字号/加粗/段落属性。"""
    from docx.oxml.ns import qn

    if style_id.lower() == "table":
        return _apply_table_style(doc, style_id, spec)
    target = _find_style(doc.styles.element, style_id)
    if target is None:
        return False
    rPr = target.find(qn("w:rPr"))
    if rPr is None:
        rPr = target.makeelement(qn("w:rPr"), {})
        _insert_style_child_in_order(target, rPr, "rPr")
    latin = spec.get("latin", spec.get("font", "Times New Roman"))
    if "font" in spec or "latin" in spec:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), latin)
        rFonts.set(qn("w:hAnsi"), latin)
        rFonts.set(qn("w:eastAsia"), spec.get("font", latin))
        rFonts.set(qn("w:cs"), latin)
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            if rFonts.get(qn(attr)) is not None:
                del rFonts.attrib[qn(attr)]
    if "size" in spec:
        half = str(int(round(spec["size"] * 2)))
        for tag in ("w:sz", "w:szCs"):
            sz = rPr.find(qn(tag))
            if sz is None:
                sz = rPr.makeelement(qn(tag), {})
                rPr.append(sz)
            sz.set(qn("w:val"), half)
    if "bold" in spec:
        b = rPr.find(qn("w:b"))
        if spec["bold"]:
            if b is None:
                b = rPr.makeelement(qn("w:b"), {})
                rPr.append(b)
            b.set(qn("w:val"), "1")
        elif b is not None:
            b.set(qn("w:val"), "0")

    # 段落属性: 首行缩进(字符)/段前段后(pt)/行距(倍数)/居中。
    pPr = target.find(qn("w:pPr"))
    need_pPr = any(k in spec for k in ("indent_chars", "before", "after", "line", "center"))
    if need_pPr:
        if pPr is None:
            pPr = target.makeelement(qn("w:pPr"), {})
            _insert_style_child_in_order(target, pPr, "pPr")
        if "indent_chars" in spec:
            ind = pPr.find(qn("w:ind"))
            if ind is None:
                ind = pPr.makeelement(qn("w:ind"), {})
                pPr.append(ind)
            ind.set(qn("w:firstLineChars"), str(spec["indent_chars"] * 100))
            ind.set(qn("w:firstLine"), str(int(spec["indent_chars"] * 240)))
        if "before" in spec or "after" in spec or "line" in spec:
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = pPr.makeelement(qn("w:spacing"), {})
                pPr.append(spacing)
            if "before" in spec:
                spacing.set(qn("w:before"), str(int(spec["before"] * 20)))
                spacing.set(qn("w:beforeLines"), "0")
            if "after" in spec:
                spacing.set(qn("w:after"), str(int(spec["after"] * 20)))
                spacing.set(qn("w:afterLines"), "0")
            if "line" in spec:
                spacing.set(qn("w:line"), str(int(spec["line"] * 240)))
                spacing.set(qn("w:lineRule"), "auto")
        if spec.get("center"):
            jc = pPr.find(qn("w:jc"))
            if jc is None:
                jc = pPr.makeelement(qn("w:jc"), {})
                pPr.append(jc)
            jc.set(qn("w:val"), "center")
    return True


def make_template(dst: Path, spec: dict[str, Any], pandoc: str | None = None,
                 base_docx: Path | None = None) -> int:
    """生成 reference.docx：pandoc 默认模板 → 按 spec 改样式。返回应用样式数。

    base_docx 提供已提取的 pandoc 默认 reference.docx 时跳过提取(本机无
    pandoc 的开发/测试场景); 生产镜像内 pandoc 存在, 自动提取。"""
    if not dst.parent.is_dir():
        dst.parent.mkdir(parents=True, exist_ok=True)
    tmpdir: Path | None = None
    if base_docx is not None and base_docx.is_file():
        base = base_docx
    else:
        pandoc = pandoc or shutil.which("pandoc")
        if not pandoc:
            print("[docx_utils] pandoc not found; cannot generate reference.docx "
                  "(镜像已预装; loopback 开发需自行安装)", file=sys.stderr)
            return 1
        # 审查: mkdtemp 目录必须清理(此前只在成功路径 unlink 文件, 目录与
        # 失败路径文件都泄漏在 /tmp)。
        tmpdir = Path(tempfile.mkdtemp(prefix="docx_utils_"))
        base = tmpdir / "reference.docx"
        try:
            # pandoc 输出为二进制 docx, 必须 -o 落文件(不能经 stdout 文本编码)。
            proc = subprocess.run([pandoc, "-o", str(base), "--print-default-data-file", "reference.docx"],
                                  capture_output=True, text=True, timeout=60)
            if proc.returncode != 0 or not base.is_file():
                print(f"[docx_utils] pandoc default reference extraction failed: {proc.stderr[:300]}", file=sys.stderr)
                shutil.rmtree(tmpdir, ignore_errors=True)
                return 1
        except subprocess.TimeoutExpired:
            print("[docx_utils] pandoc timed out", file=sys.stderr)
            shutil.rmtree(tmpdir, ignore_errors=True)
            return 1
        except OSError as exc:
            print(f"[docx_utils] pandoc execution failed: {exc}", file=sys.stderr)
            shutil.rmtree(tmpdir, ignore_errors=True)
            return 1
    try:
        from docx import Document
    except ImportError:
        print("[docx_utils] python-docx not installed", file=sys.stderr)
        if tmpdir is not None:
            shutil.rmtree(tmpdir, ignore_errors=True)
        return 1
    try:
        doc = Document(str(base))
    except Exception as exc:
        print(f"[docx_utils] 无法打开基础模板: {exc}", file=sys.stderr)
        if tmpdir is not None:
            shutil.rmtree(tmpdir, ignore_errors=True)
        return 1
    applied = 0
    missing: list[str] = []
    for style_id, s in spec.items():
        if _apply_style(doc, style_id, s):
            applied += 1
        elif style_id in AUTO_ENSURE_STYLES:
            # 语义键覆盖的样式(如 SourceCode)默认模板缺失时自动补建再应用。
            _ensure_style(doc, style_id)
            if _apply_style(doc, style_id, s):
                applied += 1
            else:
                missing.append(style_id)
        else:
            # 直写 styleId 缺失: 显式告警(拼写错误不静默造样式)。
            missing.append(style_id)
    doc.save(str(dst))
    if base_docx is None and tmpdir is not None:
        shutil.rmtree(tmpdir, ignore_errors=True)
    detail = f" ({len(missing)} 个未应用: {', '.join(missing)})" if missing else ""
    print(f"[docx_utils] ✅ 模板已生成: {dst} ({applied}/{len(spec)} 个样式应用{detail})")
    return 0 if applied > 0 else 1


# ---------------------------------------------------------------------------
# md-to-docx：pandoc 全参数转换（模板/目录/编号/GFM/高亮/编码容错）
# ---------------------------------------------------------------------------

def _read_text_robust(path: Path) -> str:
    """编码容错: UTF-8 优先, GB18030 回退(Windows 记事本来源), 兜底 replace。"""
    data = path.read_bytes()
    for enc in ("utf-8", "gb18030"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _default_reference_docx() -> Path:
    """默认参考模板: 本脚本同目录 reference.docx(ga-runner 构建期生成)。
    独立函数便于测试注入(审查 C-4: 测试不得依赖真实仓库文件系统状态)。"""
    return Path(__file__).resolve().parent / "reference.docx"


def md_to_docx(src: Path, dst: Path, template: Path | None = None, *,
               toc: bool = False, number_sections: bool = False,
               gfm: bool = False, highlight: bool = False, title: str | None = None,
               pandoc: str | None = None) -> int:
    """md/html/txt → docx。template 默认取本脚本同目录 reference.docx。

    title 非空时追加 --metadata=title:...(pandoc 官方 docx 机制: Title 样式
    段落 + docProps 元数据, 即"封面标题")。"""
    pandoc = pandoc or shutil.which("pandoc")
    if not pandoc:
        print("[docx_utils] pandoc not found (镜像已预装; loopback 需自行安装)", file=sys.stderr)
        return 1
    if template is None:
        template = _default_reference_docx()
    if not template.is_file():
        print(f"[docx_utils] reference template missing: {template} (先 make-template 生成)", file=sys.stderr)
        return 1
    if not dst.parent.is_dir():
        dst.parent.mkdir(parents=True, exist_ok=True)
    # 编码容错: GBK 源先归一 UTF-8 临时文件。
    tmp_src = src
    try:
        text = _read_text_robust(src)
        if src.suffix.lower() in (".md", ".markdown", ".html", ".htm", ".txt"):
            # 2026-08-15 审查 A-1: txt 内容命中 markdown 启发式时临时文件
            # 必须用 .md 扩展——pandoc 按扩展名推断格式, 保留 .txt 会把
            # md 当纯文本解析(标记被剥除且无排版)。
            tmp_ext = src.suffix.lower()
            if tmp_ext == ".txt" and _looks_like_markdown(text):
                tmp_ext = ".md"
            tmp_src = dst.parent / f".{dst.name}.src{tmp_ext}"
            tmp_src.write_text(text, encoding="utf-8")
    except OSError as exc:
        print(f"[docx_utils] read source failed: {exc}", file=sys.stderr)
        return 1
    cmd = [pandoc, str(tmp_src), "-o", str(dst), f"--reference-doc={template}"]
    if title:
        cmd.append(f"--metadata=title:{title}")
    if gfm:
        cmd.append("-f")
        cmd.append("gfm")
    if toc:
        cmd.append("--toc")
    if number_sections:
        cmd.append("--number-sections")
    if highlight:
        cmd.append("--highlight-style=tango")
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    except subprocess.TimeoutExpired:
        print("[docx_utils] pandoc timed out", file=sys.stderr)
        return 1
    finally:
        if tmp_src is not src:
            try:
                tmp_src.unlink(missing_ok=True)
            except OSError:
                pass
    if proc.returncode != 0:
        print(f"[docx_utils] pandoc failed: {(proc.stderr or proc.stdout or '')[:500]}", file=sys.stderr)
        return 1
    if not dst.is_file() or dst.stat().st_size == 0:
        print("[docx_utils] pandoc produced no output", file=sys.stderr)
        return 1
    print(f"[docx_utils] ✅ 已生成: {dst}")
    return 0


# ---------------------------------------------------------------------------
# verify：转换后重读验证（闭环，不靠模型自觉）
# ---------------------------------------------------------------------------

def _normalize_title_text(text: str) -> str:
    """标题对比归一(pandoc --metadata=title 按 markdown+smart 解析, 产物 Title
    段落文本与参数有可预期差异): 去 markdown 标记字符, 统一弯引号/省略号。"""
    out = re.sub(r"[#*_`~]", "", text)
    out = out.replace("\u201c", "\"").replace("\u201d", "\"")
    out = out.replace("\u2018", "'").replace("\u2019", "'")
    out = out.replace("\u2026", "...").replace("\u2014", "-")
    return out.strip()


def verify_docx(path: Path, expect_paragraphs: int | None = None,
                expect_tables: int | None = None,
                expect_title: str | None = None) -> int:
    """重读验证: 段落数下限/表格数精确/Title 样式段落(含标题文本)。

    expect_title 非空时要求存在 Title 样式段落且文本包含该字符串——对应
    pandoc --metadata=title 的产物契约(封面标题)。"""
    try:
        from docx import Document
    except ImportError:
        print("[docx_utils] python-docx not installed", file=sys.stderr)
        return 1
    try:
        doc = Document(str(path))
    except Exception as exc:
        print(f"[docx_utils] ❌ 无法打开 docx: {exc}", file=sys.stderr)
        return 1
    paras = len(doc.paragraphs)
    tables = len(doc.tables)
    headings = [p.text for p in doc.paragraphs if p.style.name.lower().startswith("heading")]
    ok = True
    # 空文档无条件失败(审查 N6: 与工具内置 _verify_docx 的段落+表格判空对齐——
    # 无任何 expect 参数时空文档也必须暴露, 不能"验证通过")。
    if paras + tables < 1:
        print(f"[docx_utils] ❌ 空文档: 段落={paras} 表格={tables}", file=sys.stderr)
        ok = False
    if expect_paragraphs is not None and paras < expect_paragraphs:
        print(f"[docx_utils] ❌ 段落数 {paras} < 期望 {expect_paragraphs}", file=sys.stderr)
        ok = False
    if expect_tables is not None and tables != expect_tables:
        print(f"[docx_utils] ❌ 表格数 {tables} != 期望 {expect_tables}", file=sys.stderr)
        ok = False
    if expect_title:
        want = _normalize_title_text(expect_title)
        titles = [_normalize_title_text(p.text)
                  for p in doc.paragraphs if p.style.name.lower() == "title"]
        if not want or not titles or not any(want in t for t in titles):
            print(f"[docx_utils] ❌ Title 样式段落缺失或未包含 {expect_title!r}", file=sys.stderr)
            ok = False
    # 审查 N1: 失败时绝不打印 ✅(此前无条件打印"验证通过"再返回 1,
    # 模型消费 stdout 会被误导)——成功/失败各打印明确状态。
    if ok:
        print(f"[docx_utils] ✅ 验证通过: {path} 段落={paras} 表格={tables} 标题={len(headings)}")
        if headings:
            print(f"[docx_utils]   标题层级: {', '.join(headings[:10])}")
    else:
        print(f"[docx_utils] ❌ 验证未通过: {path} 段落={paras} 表格={tables} 标题={len(headings)}", file=sys.stderr)
    return 0 if ok else 1


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        prog="docx_utils.py",
        description="文档转换工具脚本（pandoc 优先；模板/转换/验证一体）。",
    )
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_mk = sub.add_parser("make-template", help="按规格/预设生成 pandoc --reference-doc 模板")
    p_mk.add_argument("-o", "--output", required=True, help="输出模板路径(如 /tmp/ref.docx)")
    p_mk.add_argument("--preset", choices=("cn",), help="内置中文排版预设(社区标准)")
    p_mk.add_argument("--base", help="pandoc 默认 reference.docx(已提取, 跳过提取; 本机无 pandoc 时用)")
    p_mk.add_argument("--spec", help="样式规格 JSON，如 {\"heading1\": {\"font\": \"黑体\", \"size\": \"四号\"}}; 键支持 heading1..6/title/body/firstparagraph/bodytext/table/compact/blocktext/sourcecode 或 styleId(body/bodytext 会展开到 pandoc 正文实际使用的 FirstParagraph/BodyText)")
    p_mk.add_argument("--print-default", action="store_true", help="仅导出 pandoc 默认模板不改样式")
    p_mk.set_defaults(func=cmd_make_template)

    p_cv = sub.add_parser("md-to-docx", help="md/html/txt → docx(pandoc + 模板)")
    p_cv.add_argument("src", help="输入文件")
    p_cv.add_argument("dst", help="输出 .docx")
    p_cv.add_argument("--template", help="reference.docx 模板(默认同目录 reference.docx)")
    p_cv.add_argument("--toc", action="store_true", help="自动目录")
    p_cv.add_argument("--number-sections", action="store_true", help="章节编号")
    p_cv.add_argument("--gfm", action="store_true", help="GitHub Flavored Markdown(表格/任务列表)")
    p_cv.add_argument("--highlight", action="store_true", help="代码高亮(tango)")
    p_cv.add_argument("--title", help="封面标题(传 --metadata=title, 生成 Title 样式段落 + docProps)")
    p_cv.set_defaults(func=cmd_md_to_docx)

    p_vf = sub.add_parser("verify", help="重读验证 docx(段落/表格/标题)")
    p_vf.add_argument("path", help=".docx 文件")
    p_vf.add_argument("--expect-paragraphs", type=int, default=None)
    p_vf.add_argument("--expect-tables", type=int, default=None)
    p_vf.add_argument("--expect-title", default=None, help="要求存在 Title 样式段落且含该文本")
    p_vf.set_defaults(func=cmd_verify)

    args = parser.parse_args(argv)
    return args.func(args)


def cmd_make_template(args) -> int:
    if args.print_default:
        pandoc = shutil.which("pandoc")
        if not pandoc:
            print("pandoc not found", file=sys.stderr)
            return 1
        proc = subprocess.run([pandoc, "-o", args.output, "--print-default-data-file", "reference.docx"],
                              capture_output=True, text=True, timeout=60)
        if proc.returncode != 0:
            print(proc.stderr[:300], file=sys.stderr)
            return 1
        print(f"✅ 默认模板: {args.output}")
        return 0
    spec: dict[str, Any] = {}
    if args.preset == "cn":
        spec.update(PRESET_CN)
    if args.spec:
        try:
            user_spec = json.loads(args.spec)
        except ValueError as exc:
            print(f"spec JSON 解析失败: {exc}", file=sys.stderr)
            return 1
        spec.update(_normalize_spec(user_spec))
    if not spec:
        print("请提供 --preset cn 或 --spec", file=sys.stderr)
        return 1
    return make_template(Path(args.output), spec,
                         base_docx=Path(args.base) if args.base else None)


def cmd_md_to_docx(args) -> int:
    return md_to_docx(Path(args.src), Path(args.dst), Path(args.template) if args.template else None,
                      toc=args.toc, number_sections=args.number_sections,
                      gfm=args.gfm, highlight=args.highlight, title=args.title)


def cmd_verify(args) -> int:
    return verify_docx(Path(args.path), args.expect_paragraphs, args.expect_tables,
                       args.expect_title)


if __name__ == "__main__":
    sys.exit(main())
