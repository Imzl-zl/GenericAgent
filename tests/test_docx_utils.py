"""assets/docx_utils.py 单测（2026-08-15）。

CI/本机无 pandoc：make-template 用 --base 跳过提取；md-to-docx mock 子进程。
覆盖：预设模板样式应用（实测样式名/字体/缩进）、--spec 规格映射（中文叫法
字号）、md-to-docx 参数组装与编码容错、verify 验证闭环。
"""
import subprocess
import sys
import zipfile
from pathlib import Path

import pytest

ASSETS = Path(__file__).resolve().parent.parent / "assets"
sys.path.insert(0, str(ASSETS))

import docx_utils  # noqa: E402


@pytest.fixture
def base_reference(tmp_path: Path) -> Path:
    """pandoc 默认模板的替代品: 用 python-docx 生成含核心样式的 docx。"""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    styles = doc.styles.element
    # 模拟 pandoc reference.docx 的核心样式骨架(默认模板有 49 个, 测试子集)。
    for sid in ("Normal", "BodyText", "FirstParagraph", "Compact", "Title",
                "Heading1", "Heading2", "Heading3", "BlockText"):
        st = styles.makeelement(qn("w:style"), {})
        st.set(qn("w:type"), "paragraph")
        st.set(qn("w:styleId"), sid)
        name = st.makeelement(qn("w:name"), {})
        name.set(qn("w:val"), sid)
        st.append(name)
        styles.append(st)
    path = tmp_path / "base-ref.docx"
    doc.save(str(path))
    return path


def test_make_template_preset_cn_applies_styles(base_reference, tmp_path):
    """预设 cn: 正文 FirstParagraph 宋体小四+首行缩进+1.5 倍行距; 标题黑体加粗。"""
    dst = tmp_path / "ref.docx"
    rc = docx_utils.make_template(dst, docx_utils.PRESET_CN, base_docx=base_reference)
    assert rc == 0
    with zipfile.ZipFile(dst) as zf:
        styles = zf.read("word/styles.xml").decode("utf-8")
    assert 'w:styleId="FirstParagraph"' in styles
    assert 'w:eastAsia="宋体"' in styles
    assert 'w:val="24"' in styles, "小四=12pt=24 half-points"
    assert 'w:firstLineChars="200"' in styles, "首行缩进 2 字符"
    assert 'w:line="360"' in styles, "1.5 倍行距 = 360 twips(auto)"
    assert 'w:eastAsia="黑体"' in styles, "标题黑体"
    assert 'w:val="36"' in styles, "Heading1 小二=18pt=36 half-points"


def test_make_template_spec_semantic_keys_and_cn_size(base_reference, tmp_path):
    """--spec 语义键 + 中文叫法字号: 用户"一级黑体四号、正文宋体小四"。"""
    spec = docx_utils._normalize_spec({
        "heading1": {"font": "黑体", "size": "四号"},
        "body": {"font": "宋体", "size": "小四"},
        "heading9": {"font": "黑体"},  # 未知键忽略
    })
    assert "Heading1" in spec and spec["Heading1"]["size"] == 14.0
    # 审查 A-2: body 语义键必须展开到 pandoc 正文实际使用的样式
    # (实测: 正文段落 = FirstParagraph/BodyText, Normal 不被引用)。
    for sid in ("Normal", "BodyText", "FirstParagraph"):
        assert sid in spec and spec[sid]["size"] == 12.0, \
            f"{sid} missing from body expansion: {sorted(spec)}"
    assert "Heading9" not in spec
    dst = tmp_path / "spec.docx"
    rc = docx_utils.make_template(dst, spec, base_docx=base_reference)
    assert rc == 0
    with zipfile.ZipFile(dst) as zf:
        styles = zf.read("word/styles.xml").decode("utf-8")
    assert 'w:val="28"' in styles, "四号=14pt=28 half-points"
    assert 'w:styleId="FirstParagraph"' in styles
    assert 'w:styleId="BodyText"' in styles


def test_make_template_ensures_missing_semantic_style(base_reference, tmp_path):
    """语义键覆盖的样式在默认模板缺失时(如 SourceCode)必须自动补建再应用,
    不能静默丢弃(审查 C-1: 实测 pandoc 默认模板 12/13, 缺 SourceCode)。"""
    dst = tmp_path / "ref.docx"
    rc = docx_utils.make_template(dst, docx_utils.PRESET_CN, base_docx=base_reference)
    assert rc == 0
    with zipfile.ZipFile(dst) as zf:
        styles = zf.read("word/styles.xml").decode("utf-8")
    assert 'w:styleId="SourceCode"' in styles, "SourceCode 缺失时应自动补建"
    assert 'w:ascii="Consolas"' in styles, "补建后应应用 Consolas 字体"


def test_make_template_missing_unknown_style_reported(tmp_path, base_reference):
    """直写未知 styleId 缺失时显式告警(不静默造样式, 拼写错误可见)。"""
    dst = tmp_path / "ref.docx"
    rc = docx_utils.make_template(dst, {"Heading 1": {"font": "黑体"}},
                                  base_docx=base_reference)
    assert rc == 1, "unknown style must not silently create a template"
    with zipfile.ZipFile(dst) as zf:
        styles = zf.read("word/styles.xml").decode("utf-8")
    assert 'w:styleId="Heading 1"' not in styles


def test_make_template_missing_pandoc_without_base_fails():
    """无 pandoc 且无 --base: 必须显式失败(不产出残缺模板)。"""
    rc = docx_utils.make_template(Path("/tmp/never.docx"), {"Normal": {"font": "宋体"}},
                                  pandoc="/nonexistent/pandoc")
    assert rc == 1


def test_md_to_docx_builds_pandoc_command_and_handles_gbk(tmp_path, monkeypatch):
    """md-to-docx: 参数组装(--toc/--number-sections/--gfm/模板)、GBK 编码归一。"""
    src = tmp_path / "输入.txt"
    src.write_bytes("工作计划：今天评审。\n".encode("gb18030"))
    dst = tmp_path / "out.docx"
    template = tmp_path / "ref.docx"
    template.write_bytes(b"PK-fake")
    calls = []
    tmp_contents = []

    def fake_run(cmd, capture_output=True, text=True, timeout=120):
        calls.append(cmd)
        tmp_src_path = next((c for c in cmd if ".src.txt" in c), None)
        if tmp_src_path:
            tmp_contents.append(Path(tmp_src_path).read_text(encoding="utf-8"))
        dst.write_bytes(b"PK\x03\x04fake-docx")
        return subprocess.CompletedProcess(cmd, 0, stdout="", stderr="")

    monkeypatch.setattr(docx_utils.subprocess, "run", fake_run)
    rc = docx_utils.md_to_docx(src, dst, template, toc=True, number_sections=True,
                               gfm=True, highlight=True, pandoc="/usr/bin/pandoc")
    assert rc == 0
    assert calls, "pandoc must be invoked"
    cmd = calls[0]
    assert cmd[0] == "/usr/bin/pandoc"
    assert f"--reference-doc={template}" in cmd
    assert "--toc" in cmd and "--number-sections" in cmd and "--highlight-style=tango" in cmd
    assert cmd[cmd.index("-f") + 1] == "gfm"
    # GBK 源: 临时源文件为 UTF-8 且已清理。
    assert tmp_contents == ["工作计划：今天评审。\n"], "GBK 源必须归一为 UTF-8"
    assert all(not Path(c).exists() for c in cmd if ".src.txt" in c), "temp source must be cleaned up"


def test_md_to_docx_txt_with_markdown_uses_md_extension(tmp_path, monkeypatch):
    """审查 A-1: txt 内容命中 markdown 启发式时临时源必须用 .md 扩展——
    pandoc 按扩展名推断格式, 保留 .txt 会把 md 当纯文本解析(标记被剥除
    且无排版, 实测)。"""
    src = tmp_path / "note.txt"
    src.write_text("# 标题\n\n- 项一\n- 项二\n", encoding="utf-8")
    dst = tmp_path / "out.docx"
    template = tmp_path / "ref.docx"
    template.write_bytes(b"PK-fake")
    calls = []

    def fake_run(cmd, capture_output=True, text=True, timeout=120):
        calls.append(cmd)
        dst.write_bytes(b"PK\x03\x04fake-docx")
        return subprocess.CompletedProcess(cmd, 0, stdout="", stderr="")

    monkeypatch.setattr(docx_utils.subprocess, "run", fake_run)
    rc = docx_utils.md_to_docx(src, dst, template, pandoc="/usr/bin/pandoc")
    assert rc == 0
    assert calls
    tmp_src = next(c for c in calls[0] if ".src" in c)
    assert tmp_src.endswith(".src.md"), \
        f"markdown-looking txt must use .md ext, got {tmp_src}"
    assert not Path(tmp_src).exists(), "temp source must be cleaned up"


def test_md_to_docx_plain_txt_keeps_txt_extension(tmp_path, monkeypatch):
    """纯文本 txt(无 markdown 结构)保留 .txt 扩展(按纯文本段落转换)。"""
    src = tmp_path / "plain.txt"
    src.write_text("第一行\n第二行\n", encoding="utf-8")
    dst = tmp_path / "out.docx"
    template = tmp_path / "ref.docx"
    template.write_bytes(b"PK-fake")
    calls = []

    def fake_run(cmd, capture_output=True, text=True, timeout=120):
        calls.append(cmd)
        dst.write_bytes(b"PK\x03\x04fake-docx")
        return subprocess.CompletedProcess(cmd, 0, stdout="", stderr="")

    monkeypatch.setattr(docx_utils.subprocess, "run", fake_run)
    rc = docx_utils.md_to_docx(src, dst, template, pandoc="/usr/bin/pandoc")
    assert rc == 0
    tmp_src = next(c for c in calls[0] if ".src" in c)
    assert tmp_src.endswith(".src.txt"), f"plain txt keeps .txt ext, got {tmp_src}"


def test_md_to_docx_missing_template_fails(tmp_path, monkeypatch):
    src = tmp_path / "a.md"
    src.write_text("# x\n", encoding="utf-8")
    # 审查 C-4: 默认模板路径可注入, 测试不依赖仓库真实文件系统状态
    # (worker 集成测试的占位 fixture 会临时创建 assets/reference.docx,
    # 并行 pytest 进程下依赖真实路径的断言会 flaky)。
    monkeypatch.setattr(docx_utils, "_default_reference_docx", lambda: tmp_path / "missing.docx")
    rc = docx_utils.md_to_docx(src, tmp_path / "b.docx", None,
                               pandoc="/usr/bin/pandoc")
    assert rc == 1, "missing template must fail explicitly (default 模板未生成时)"


def test_verify_docx_reports_and_checks(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_heading("标题", level=1)
    doc.add_paragraph("正文")
    doc.add_table(rows=2, cols=2)
    path = tmp_path / "v.docx"
    doc.save(str(path))
    assert docx_utils.verify_docx(path) == 0
    assert docx_utils.verify_docx(path, expect_tables=1) == 0
    assert docx_utils.verify_docx(path, expect_tables=2) == 1
    assert docx_utils.verify_docx(tmp_path / "missing.docx") == 1


def test_cn_font_size_table():
    assert docx_utils.CN_FONT_SIZE_PT["四号"] == 14
    assert docx_utils.CN_FONT_SIZE_PT["小四"] == 12
    assert docx_utils._resolve_pt("三号") == 16
    assert docx_utils._resolve_pt(15) == 15
    assert docx_utils._resolve_pt("超大") is None
