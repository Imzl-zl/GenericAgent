"""双轨 helper 漂移守卫(审查 M4): assets/docx_utils.py(独立脚本, 镜像内
/ga/legacy/assets/)与 ga_worker.legacy_instrument.py(worker 进程内)各有一份
_looks_like_markdown / _read_text_robust / _normalize_title_text 副本——两个
运行时上下文(脚本 vs 模块)结构性无法合并, 收敛义务由本测试落地: 行为不一致
即测试失败, 防静默漂移(改一处忘另一处)。"""

from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parents[4]  # .../tests/unit → 仓库根


@pytest.fixture(scope="module")
def assets_docx_utils():
    """从仓库 assets/ 加载 docx_utils.py(不经 sys.path 安装)。"""
    path = REPO_ROOT / "assets" / "docx_utils.py"
    spec = importlib.util.spec_from_file_location("docx_utils_under_drift_test", path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def test_looks_like_markdown_equivalent(assets_docx_utils):
    from ga_worker import legacy_instrument

    corpus = [
        "", "   ", "纯文本一行", "第一行\n第二行\n",
        "# 标题\n\n正文", "## 二级\n- 列表", "- 项一\n- 项二",
        "| a | b |\n|---|---|\n| 1 | 2 |",
        "文本 **加粗** 文本", "```python\nprint(1)\n```",
        "> 引用\n\n普通", "1. 有序\n2. 有序", "代码 `inline` 混排",
        "http://example.com 链接", "表格前说明\n\n| x |\n|---|\n| 1 |",
        "#标题无空格不识别", "### 三级标题\n正文",
    ]
    for sample in corpus:
        a = assets_docx_utils._looks_like_markdown(sample)
        b = legacy_instrument._looks_like_markdown(sample)
        assert a == b, f"_looks_like_markdown drift on {sample!r}: assets={a} worker={b}"


def test_read_text_robust_equivalent(assets_docx_utils, tmp_path: Path):
    from ga_worker import legacy_instrument

    # UTF-8 / GB18030(GBK 超集) / 非法字节兜底 三类输入。
    cases = [
        ("utf8.txt", "中文内容 UTF-8".encode("utf-8")),
        ("gbk.txt", "中文内容 GBK 编码".encode("gb18030")),
        ("broken.txt", b"\xff\xfe broken \x80\x81"),
        ("ascii.txt", b"plain ascii"),
    ]
    for name, data in cases:
        p = tmp_path / name
        p.write_bytes(data)
        a = assets_docx_utils._read_text_robust(p)
        b = legacy_instrument._read_text_robust(p)
        assert a == b, f"_read_text_robust drift on {name}"


def test_normalize_title_text_equivalent(assets_docx_utils):
    from ga_worker import legacy_instrument

    corpus = [
        "普通标题", "计划：“Q3”报告", "计划：\"Q3\"报告", "标题 '单引号' 内容",
        "*强调*标题", "**加粗**标题", "带`反引号`标题", "省略号...测试",
        "破折号——测试", " 前后空白  ", "#Markdown前缀", "~~删除线~~标题",
        "中文，标点。测试", "A—B 与 A-B 混合",
    ]
    for sample in corpus:
        a = assets_docx_utils._normalize_title_text(sample)
        b = legacy_instrument._normalize_title_text(sample)
        assert a == b, f"_normalize_title_text drift on {sample!r}: assets={a!r} worker={b!r}"


def test_verify_title_assertion_semantics_equivalent(assets_docx_utils, tmp_path: Path):
    """工具 _verify_docx 与脚本 verify_docx 的 Title 断言语义一致(归一+子串)。"""
    from docx import Document

    from ga_worker import legacy_instrument

    doc = Document()
    doc.add_paragraph("计划：“Q3”报告", style="Title")
    doc.add_paragraph("正文")
    p = tmp_path / "v.docx"
    doc.save(str(p))

    ok, _ = legacy_instrument._verify_docx(p, expect_title='计划："Q3"报告')
    assert ok, "worker verify must accept smart-quoted title"
    assert assets_docx_utils.verify_docx(p, expect_title='计划："Q3"报告') == 0

    ok, _ = legacy_instrument._verify_docx(p, expect_title="不存在的标题")
    assert not ok, "worker verify must reject mismatched title"
    assert assets_docx_utils.verify_docx(p, expect_title="不存在的标题") == 1
