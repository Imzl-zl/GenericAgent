from __future__ import annotations

import types
from pathlib import Path

from genericagent.worker.v1 import worker_pb2
from ga_worker.legacy_instrument import (
    apply_tool_policy,
    install_dispatch_guard,
    install_export_docx_tool,
    install_global_mcp_tools,
    ensure_session_dirs,
    restore_tool_schema,
)
from ga_worker.limits import ToolPolicy
from ga_worker.session_files import append_missing_file_markers, ensure_session_sandbox, normalize_output_name, resolve_under_root
from ga_worker.state import PendingTask, TaskRunState
from ga_worker.task_terminal import emit_cancel_or_timeout_terminal, emit_final_terminal


def test_resolve_under_root_rejects_escape(tmp_path: Path):
    root = ensure_session_sandbox(tmp_path, "personal:1")
    resolved = resolve_under_root(root, "outputs/resume.docx")
    assert resolved == root / "outputs" / "resume.docx"

    try:
        resolve_under_root(root, "../escape.txt")
    except ValueError as exc:
        assert "escapes session sandbox" in str(exc)
    else:
        raise AssertionError("expected sandbox escape rejection")


def test_normalize_output_name_preserves_outputs_prefix():
    assert normalize_output_name("outputs/resume.docx") == "outputs/resume.docx"
    assert normalize_output_name("resume") == "outputs/resume.docx"
    assert normalize_output_name("outputs/subdir/resume") == "outputs/subdir/resume.docx"



def test_append_missing_file_markers_idempotent():
    body = "搞定了"
    with_marker = append_missing_file_markers(body, ["outputs/a.docx", "outputs/b.docx"])
    assert "[FILE:outputs/a.docx]" in with_marker
    assert "[FILE:outputs/b.docx]" in with_marker
    again = append_missing_file_markers(with_marker, ["outputs/a.docx"])
    assert again.count("[FILE:outputs/a.docx]") == 1



def test_emit_final_terminal_appends_generated_output_markers():
    recorded: dict[str, str] = {}

    class FakeAdapter:
        def __init__(self):
            self._session = types.SimpleNamespace(generated_output_files=["outputs/demo.docx"])

        def _terminal(self, task_id, status, *, user_message="", error_code=None, result_body=None):
            return worker_pb2.Terminal(task_id=task_id, status=status, user_message=user_message)

        def _record_completed(self, task, term, result_body, display_history, agent):
            recorded["body"] = result_body

    task = worker_pb2.TaskEnvelope(task_id="t1", session_key="personal:1")
    state = TaskRunState(
        pending=PendingTask(task_id="t1", request=worker_pb2.ExecuteTaskRequest()),
        agent=object(),
        final_body="搞定",
    )
    events = list(emit_final_terminal(FakeAdapter(), task, state))
    assert len(events) == 1
    assert "[FILE:outputs/demo.docx]" in state.final_body
    assert "[FILE:outputs/demo.docx]" in recorded["body"]





def test_export_docx_tool_injected_and_generates_real_docx(tmp_path: Path, monkeypatch):
    """round11 审查 I5: export_docx 必须真实注入(policy 允许 + 提示引导),
    且能生成可解析的 .docx 文件, 路径逃逸被拒绝。"""
    import sys
    import zipfile

    class StepOutcome:
        def __init__(self, data, next_prompt=None, should_exit=False):
            self.data = data
            self.next_prompt = next_prompt
            self.should_exit = should_exit

    monkeypatch.setitem(sys.modules, "agent_loop", types.SimpleNamespace(StepOutcome=StepOutcome))

    class FakeHandler:
        def __init__(self, *args, **kwargs):
            self.cwd = "./temp"

        def _get_abs_path(self, path):
            import os as _os
            return _os.path.abspath(_os.path.join(self.cwd, path))

    ga_mod = types.SimpleNamespace(GenericAgentHandler=FakeHandler)
    agentmain_mod = types.SimpleNamespace(TOOLS_SCHEMA=[
        {"type": "function", "function": {"name": "file_read"}},
    ])
    session = types.SimpleNamespace(
        overlay_dir=tmp_path / "runtime" / "s-1" / "legacy-overlay",
        session_key="personal:1",
        agent=types.SimpleNamespace(handler=None),
        generated_output_files=[],
    )
    session.overlay_dir.mkdir(parents=True, exist_ok=True)

    unwrap = install_export_docx_tool(session, {"ga": ga_mod, "agentmain": agentmain_mod})
    previous = apply_tool_policy(
        ToolPolicy(version="foundation.session-files.v1", allowed_tools=frozenset({"file_read", "export_docx"})),
        {"ga": ga_mod, "agentmain": agentmain_mod},
    )
    try:
        names = [t["function"]["name"] for t in agentmain_mod.TOOLS_SCHEMA]
        assert "export_docx" in names

        handler = ga_mod.GenericAgentHandler(None)
        out = list(handler.do_export_docx(
            {"path": "outputs/resume.docx", "title": "简历", "content": "第一段\n第二段"},
            "",
        ))
        assert len(out) >= 1
        target = tmp_path / "runtime" / "session_files" / "personal:1" / "outputs" / "resume.docx"
        # session_sandbox_root 使用 session_key digest 布局(无 GA_WORKSPACE_TEMP)。
        from ga_worker.session_files import session_key_digest
        target = tmp_path / "runtime" / "session_files" / session_key_digest("personal:1") / "outputs" / "resume.docx"
        assert target.exists(), f"docx not generated at {target}"
        with zipfile.ZipFile(target) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert "简历" in xml
        assert "第一段" in xml
        assert "outputs/resume.docx" in session.generated_output_files

        # 路径逃逸拒绝: 不生成文件且返回错误 outcome。
        out_escape = list(handler.do_export_docx(
            {"path": "../../escape.docx", "content": "x"},
            "",
        ))
        assert not (tmp_path / "escape.docx").exists()
        assert len(out_escape) >= 1
    finally:
        restore_tool_schema(previous, {"ga": ga_mod, "agentmain": agentmain_mod})
        unwrap()

    # unwrap 后工具不再存在。
    assert not hasattr(ga_mod.GenericAgentHandler, "do_export_docx")


def test_export_docx_tool_reads_source_path(tmp_path: Path, monkeypatch):
    import sys

    class StepOutcome:
        def __init__(self, data, next_prompt=None, should_exit=False):
            self.data = data
            self.next_prompt = next_prompt
            self.should_exit = should_exit

    monkeypatch.setitem(sys.modules, "agent_loop", types.SimpleNamespace(StepOutcome=StepOutcome))

    class FakeHandler:
        def __init__(self, *args, **kwargs):
            self.cwd = "./temp"

    ga_mod = types.SimpleNamespace(GenericAgentHandler=FakeHandler)
    agentmain_mod = types.SimpleNamespace(TOOLS_SCHEMA=[])
    session = types.SimpleNamespace(
        overlay_dir=tmp_path / "runtime" / "s-1" / "legacy-overlay",
        session_key="personal:1",
        generated_output_files=[],
    )
    session.overlay_dir.mkdir(parents=True, exist_ok=True)

    from ga_worker.session_files import session_key_digest, session_sandbox_root
    root = session_sandbox_root(tmp_path / "runtime", "personal:1")
    src = root / "temp_notes.txt"
    src.parent.mkdir(parents=True, exist_ok=True)
    src.write_text("源文件内容", encoding="utf-8")

    unwrap = install_export_docx_tool(session, {"ga": ga_mod, "agentmain": agentmain_mod})
    try:
        handler = ga_mod.GenericAgentHandler(None)
        list(handler.do_export_docx(
            {"path": "outputs/from_source.docx", "source_path": "temp_notes.txt"},
            "",
        ))
        target = root / "outputs" / "from_source.docx"
        assert target.exists()
    finally:
        unwrap()


def test_global_mcp_tool_runs_when_tenant_policy_allows_it(tmp_path: Path, monkeypatch):
    import sys

    class StepOutcome:
        def __init__(self, data, next_prompt=None, should_exit=False):
            self.data = data
            self.next_prompt = next_prompt
            self.should_exit = should_exit

    monkeypatch.setitem(sys.modules, "agent_loop", types.SimpleNamespace(StepOutcome=StepOutcome))

    class FakeHandler:
        def __init__(self, *args, **kwargs):
            self.cwd = "./temp"

        def dispatch(self, tool_name, args, response, index=0, tool_num=1):
            method = getattr(self, f"do_{tool_name}", None)
            if method is None:
                yield f"unknown:{tool_name}\n"
                return StepOutcome(None)
            return (yield from method(args, response))

    class FakeClient:
        def call_tool(self, name, arguments):
            assert name == "web_search"
            assert arguments == {"query": "GA"}
            return "search result"

    ga_mod = types.SimpleNamespace(GenericAgentHandler=FakeHandler)
    agentmain_mod = types.SimpleNamespace(TOOLS_SCHEMA=[
        {"type": "function", "function": {"name": "file_read"}},
    ])
    session = types.SimpleNamespace(
        overlay_dir=tmp_path / "runtime" / "s-1" / "legacy-overlay",
        session_key="personal:1",
        agent=types.SimpleNamespace(handler=None),
        generated_output_files=[],
        mcp_tools={
            "exa__web_search": {
                "client": FakeClient(),
                "tool_name": "web_search",
                "schema": {
                    "type": "function",
                    "function": {
                        "name": "exa__web_search",
                        "description": "Search",
                        "parameters": {"type": "object", "properties": {"query": {"type": "string"}}},
                    },
                },
            },
        },
    )
    session.overlay_dir.mkdir(parents=True, exist_ok=True)
    mods = {"ga": ga_mod, "agentmain": agentmain_mod, "agent_loop": sys.modules["agent_loop"]}

    sandbox_unwrap = ensure_session_dirs(session, mods)
    mcp_unwrap = install_global_mcp_tools(session, mods)
    policy = ToolPolicy(
        version="foundation.session-files.v1",
        allowed_tools=frozenset({"file_read", "export_docx", "exa__web_search"}),
    )
    previous = apply_tool_policy(policy, mods)
    guard_unwrap = install_dispatch_guard(policy, mods)
    try:
        names = [tool["function"]["name"] for tool in agentmain_mod.TOOLS_SCHEMA]
        assert names == ["file_read", "exa__web_search"]

        call = ga_mod.GenericAgentHandler().dispatch("exa__web_search", {"query": "GA"}, None)
        chunks = []
        try:
            while True:
                chunks.append(next(call))
        except StopIteration as stop:
            assert stop.value.data == "search result"
        assert any("search result" in chunk for chunk in chunks)
    finally:
        guard_unwrap()
        restore_tool_schema(previous, mods)
        mcp_unwrap()
        sandbox_unwrap()


def test_global_mcp_install_is_transactional_when_catalog_contains_conflict():
    class FakeHandler:
        def do_conflict__existing(self, args, response):
            return None

    ga_mod = types.SimpleNamespace(GenericAgentHandler=FakeHandler)
    agentmain_mod = types.SimpleNamespace(TOOLS_SCHEMA=[])

    def binding(name: str):
        return {
            "client": object(),
            "tool_name": name,
            "schema": {
                "type": "function",
                "function": {
                    "name": name,
                    "description": name,
                    "parameters": {"type": "object", "properties": {}},
                },
            },
        }

    session = types.SimpleNamespace(mcp_tools={
        "ok__search": binding("ok__search"),
        "conflict__existing": binding("conflict__existing"),
    })
    try:
        install_global_mcp_tools(session, {"ga": ga_mod, "agentmain": agentmain_mod})
    except ValueError as exc:
        assert "conflicts" in str(exc)
    else:
        raise AssertionError("expected MCP handler conflict")

    assert not hasattr(FakeHandler, "do_ok__search")
    assert not hasattr(agentmain_mod, "_tenant_custom_tools_schema")
    assert not hasattr(agentmain_mod, "_tenant_global_mcp_tool_names")


def test_global_mcp_tools_allow_mcp_wildcard(tmp_path: Path, monkeypatch):
    # mcp:* 通配: apply_tool_policy 与 dispatch_guard 都放行 MCP 工具,
    # 非 MCP 工具仍精确匹配(deny-by-default)。
    import sys

    class StepOutcome:
        def __init__(self, data, next_prompt=None, should_exit=False):
            self.data = data
            self.next_prompt = next_prompt
            self.should_exit = should_exit

    agent_loop_mod = types.SimpleNamespace(StepOutcome=StepOutcome)
    monkeypatch.setitem(sys.modules, "agent_loop", agent_loop_mod)

    dispatched: list[str] = []

    class FakeHandler:
        def __init__(self, *args, **kwargs):
            self.cwd = "./temp"

        def dispatch(self, tool_name, args, response, index=0, tool_num=1):
            dispatched.append(tool_name)
            method = getattr(self, f"do_{tool_name}", None)
            if method is None:
                yield f"unknown:{tool_name}\n"
                return StepOutcome(None)
            return (yield from method(args, response))

    class FakeClient:
        def call_tool(self, remote_name, public_args):
            return f"result:{remote_name}"

    ga_mod = types.SimpleNamespace(GenericAgentHandler=FakeHandler)
    agentmain_mod = types.SimpleNamespace(TOOLS_SCHEMA=[])
    session = types.SimpleNamespace(
        mcp_tools={
            "exa__web_search_exa": {
                "client": FakeClient(),
                "tool_name": "web_search_exa",
                "schema": {
                    "type": "function",
                    "function": {
                        "name": "exa__web_search_exa",
                        "description": "Search",
                        "parameters": {"type": "object", "properties": {"query": {"type": "string"}}},
                    },
                },
            },
        },
    )
    mods = {"ga": ga_mod, "agentmain": agentmain_mod, "agent_loop": sys.modules["agent_loop"]}
    mcp_unwrap = install_global_mcp_tools(session, mods)
    policy = ToolPolicy(
        version="foundation.session-files.v1",
        allowed_tools=frozenset({"file_read", "mcp:*"}),
    )
    previous = apply_tool_policy(policy, mods)
    guard_unwrap = install_dispatch_guard(policy, mods)
    try:
        names = [tool["function"]["name"] for tool in agentmain_mod.TOOLS_SCHEMA]
        assert names == ["exa__web_search_exa"], names

        handler = ga_mod.GenericAgentHandler()
        call = handler.dispatch("exa__web_search_exa", {"query": "GA"}, None)
        chunks = []
        try:
            while True:
                chunks.append(next(call))
        except StopIteration as stop:
            assert stop.value.data == "result:web_search_exa"
        assert dispatched == ["exa__web_search_exa"]
        # 非 MCP 工具(不在白名单)仍被守卫拒绝。
        denied = handler.dispatch("file_write", {}, None)
        assert "tool denied by policy" in next(denied)
    finally:
        guard_unwrap()
        restore_tool_schema(previous, mods)
        mcp_unwrap()


def test_global_mcp_tools_must_intersect_with_tenant_policy(tmp_path: Path, monkeypatch):
    import sys

    class StepOutcome:
        def __init__(self, data, next_prompt=None, should_exit=False):
            self.data = data
            self.next_prompt = next_prompt
            self.should_exit = should_exit

    agent_loop_mod = types.SimpleNamespace(StepOutcome=StepOutcome)
    monkeypatch.setitem(sys.modules, "agent_loop", agent_loop_mod)

    dispatched: list[str] = []

    class FakeHandler:
        def __init__(self, *args, **kwargs):
            self.cwd = "./temp"

        def dispatch(self, tool_name, args, response, index=0, tool_num=1):
            dispatched.append(tool_name)
            method = getattr(self, f"do_{tool_name}", None)
            if method is None:
                yield f"unknown:{tool_name}\n"
                return StepOutcome(None)
            return (yield from method(args, response))

    class FakeClient:
        def call_tool(self, name, arguments):
            assert name == "web_search"
            assert arguments == {"query": "GA"}
            return "search result"

    ga_mod = types.SimpleNamespace(GenericAgentHandler=FakeHandler)
    agentmain_mod = types.SimpleNamespace(TOOLS_SCHEMA=[
        {"type": "function", "function": {"name": "file_read"}},
        {"type": "function", "function": {"name": "update_working_checkpoint"}},
    ])
    session = types.SimpleNamespace(
        overlay_dir=tmp_path / "runtime" / "s-1" / "legacy-overlay",
        session_key="personal:1",
        agent=types.SimpleNamespace(handler=None),
        generated_output_files=[],
        mcp_tools={
            "exa__web_search": {
                "client": FakeClient(),
                "tool_name": "web_search",
                "schema": {
                    "type": "function",
                    "function": {
                        "name": "exa__web_search",
                        "description": "Search",
                        "parameters": {"type": "object", "properties": {"query": {"type": "string"}}},
                    },
                },
            },
        },
    )
    session.overlay_dir.mkdir(parents=True, exist_ok=True)
    mods = {"ga": ga_mod, "agentmain": agentmain_mod, "agent_loop": agent_loop_mod}

    sandbox_unwrap = ensure_session_dirs(session, mods)
    mcp_unwrap = install_global_mcp_tools(session, mods)
    policy = ToolPolicy(
        version="foundation.session-files.v1",
        allowed_tools=frozenset({"file_read", "export_docx"}),
    )
    previous = apply_tool_policy(policy, mods)
    guard_unwrap = install_dispatch_guard(policy, mods)
    try:
        names = [tool["function"]["name"] for tool in agentmain_mod.TOOLS_SCHEMA]
        assert names == ["file_read"]

        handler = ga_mod.GenericAgentHandler()
        denied = handler.dispatch(
            "exa__web_search",
            {"query": "GA", "_index": 0, "_tool_num": 1},
            None,
        )
        denied_chunks = []
        try:
            while True:
                denied_chunks.append(next(denied))
        except StopIteration:
            pass
        assert any("denied" in chunk for chunk in denied_chunks)
        assert dispatched == []
    finally:
        guard_unwrap()
        restore_tool_schema(previous, mods)
        mcp_unwrap()
        sandbox_unwrap()


def test_image_gen_marker_registry_registers_outputs(tmp_path: Path, monkeypatch):
    """Phase B 托管形态(I-2): image_gen 的 marker 回显兜底——包装 GA
    do_image_gen(generator 转发), 工具成功返回 [FILE:outputs/...] 后登记到
    session.generated_output_files, 供终态 append_missing_file_markers 兜底。
    包装不替换工具逻辑(原方法仍被调用)。"""
    from ga_worker.legacy_instrument import install_image_gen_marker_registry

    called = []

    def fake_do_image_gen(self, args, response):
        called.append(args)

        def gen():
            yield "[Action] image_gen: ...\n"
            return types.SimpleNamespace(
                data="[FILE:outputs/image_20260814_1.png]\n[FILE:outputs/image_20260814_2.png]",
                next_prompt="\n[FILE:outputs/image_20260814_1.png]",
                should_exit=False,
            )
        return gen()

    class FakeHandler:
        do_image_gen = fake_do_image_gen

    ga_mod = types.SimpleNamespace(GenericAgentHandler=FakeHandler)
    session = types.SimpleNamespace(generated_output_files=[])

    unwrap = install_image_gen_marker_registry(session, {"ga": ga_mod})
    handler = ga_mod.GenericAgentHandler()
    out = list(handler.do_image_gen({"prompt": "a cat", "n": 2}, None))
    assert out == ["[Action] image_gen: ...\n"], out
    # 登记: 工具返回的 marker 全部进 generated_output_files(去重)
    assert session.generated_output_files == [
        "outputs/image_20260814_1.png",
        "outputs/image_20260814_2.png",
    ]

    unwrap()
    assert ga_mod.GenericAgentHandler.do_image_gen is fake_do_image_gen


def test_image_gen_marker_registry_error_returns_no_registration(tmp_path: Path, monkeypatch):
    """失败语义: 工具返回错误文本(无 [FILE:])时零登记——模型不得谎报成功。"""
    from ga_worker.legacy_instrument import install_image_gen_marker_registry

    def fake_do_image_gen(self, args, response):
        def gen():
            yield "[Status] ❌ ...\n"
            return types.SimpleNamespace(
                data="[Error: image_gen HTTP 500]",
                next_prompt="\n",
                should_exit=False,
            )
        return gen()

    class FakeHandler:
        do_image_gen = fake_do_image_gen

    ga_mod = types.SimpleNamespace(GenericAgentHandler=FakeHandler)
    session = types.SimpleNamespace(generated_output_files=[])
    unwrap = install_image_gen_marker_registry(session, {"ga": ga_mod})
    list(ga_mod.GenericAgentHandler().do_image_gen({"prompt": "x"}, None))
    assert session.generated_output_files == []
    unwrap()


def _export_docx_harness(tmp_path: Path, monkeypatch):
    """export_docx 测试脚手架: 返回 (handler, root, session, unwrap, legacy_instrument_mod)。"""
    import sys
    import types

    class StepOutcome:
        def __init__(self, data, next_prompt=None, should_exit=False):
            self.data = data
            self.next_prompt = next_prompt
            self.should_exit = should_exit

    monkeypatch.setitem(sys.modules, "agent_loop", types.SimpleNamespace(StepOutcome=StepOutcome))

    class FakeHandler:
        def __init__(self, *args, **kwargs):
            self.cwd = "./temp"

    ga_mod = types.SimpleNamespace(GenericAgentHandler=FakeHandler)
    agentmain_mod = types.SimpleNamespace(TOOLS_SCHEMA=[])
    session = types.SimpleNamespace(
        overlay_dir=tmp_path / "runtime" / "s-1" / "legacy-overlay",
        session_key="personal:1",
        generated_output_files=[],
    )
    session.overlay_dir.mkdir(parents=True, exist_ok=True)

    from ga_worker import legacy_instrument
    from ga_worker.session_files import session_sandbox_root

    root = session_sandbox_root(tmp_path / "runtime", "personal:1")
    unwrap = legacy_instrument.install_export_docx_tool(session, {"ga": ga_mod, "agentmain": agentmain_mod})
    handler = ga_mod.GenericAgentHandler(None)
    return handler, root, session, unwrap, legacy_instrument


def _make_fake_pandoc(target: Path):
    """fake pandoc: 用 python-docx 生成一个含标题样式的 docx(模拟 pandoc 产物),
    返回可注入 _run_pandoc 的替身。"""
    from docx import Document

    def _fake(src, tgt, timeout_s=60):
        assert Path(src).is_file(), f"pandoc source missing: {src}"
        doc = Document()
        doc.add_heading("转换标题", level=1)
        doc.add_paragraph("正文段落")
        doc.save(str(tgt))
        return True, ""

    return _fake


def test_export_docx_markdown_source_routes_to_pandoc(tmp_path: Path, monkeypatch):
    """2026-08-15: md 源必须走 pandoc(不逐行裸 dump), 产物统一中文字体。"""
    handler, root, session, unwrap, mod = _export_docx_harness(tmp_path, monkeypatch)
    try:
        src = root / "notes.md"
        src.parent.mkdir(parents=True, exist_ok=True)
        src.write_text("# 标题\n\n- 列表项\n\n| a | b |\n|---|---|\n| 1 | 2 |\n", encoding="utf-8")
        calls = []
        fake = _make_fake_pandoc(root / "outputs" / "x.docx")
        monkeypatch.setattr(mod, "_run_pandoc", fake)

        out = list(handler.do_export_docx(
            {"path": "outputs/from_md.docx", "source_path": "notes.md"},
            "",
        ))
        target = root / "outputs" / "from_md.docx"
        assert target.exists()
        # 工具如实报告引擎。
        assert any("pandoc" in str(o) for o in out), f"engine not reported: {out}"
        # 中文字体统一(pandoc 默认模板混排修复)。
        import zipfile
        with zipfile.ZipFile(target) as zf:
            styles = zf.read("word/styles.xml").decode("utf-8")
        assert "微软雅黑" in styles
    finally:
        unwrap()


def test_export_docx_markdown_content_routes_to_pandoc(tmp_path: Path, monkeypatch):
    """2026-08-15: content 为 markdown 结构时走 pandoc; 纯文本不依赖 pandoc。"""
    handler, root, session, unwrap, mod = _export_docx_harness(tmp_path, monkeypatch)
    try:
        calls = []
        fake = _make_fake_pandoc(root / "outputs" / "y.docx")
        monkeypatch.setattr(mod, "_run_pandoc", lambda src, tgt, timeout_s=60: calls.append((str(src), str(tgt))) or fake(src, tgt))

        out = list(handler.do_export_docx(
            {"path": "outputs/md_content.docx", "content": "# 大标题\n\n**加粗** 正文\n\n- 项一\n- 项二\n"},
            "",
        ))
        target = root / "outputs" / "md_content.docx"
        assert target.exists(), f"docx missing: {target}"
        assert len(calls) == 1, f"pandoc must be invoked for markdown content, calls={calls}"
        assert calls[0][0].endswith(".src.md"), f"temp md source expected, got {calls[0][0]}"
        # 临时源已清理。
        assert not Path(calls[0][0]).exists()
        assert any("pandoc" in str(o) for o in out)

        # 纯文本: 不调 pandoc, 直接 python-docx。
        calls.clear()
        out2 = list(handler.do_export_docx(
            {"path": "outputs/plain.docx", "title": "普通文档", "content": "第一行\n第二行"},
            "",
        ))
        target2 = root / "outputs" / "plain.docx"
        assert target2.exists()
        assert len(calls) == 0, f"pandoc must NOT be invoked for plain text, calls={calls}"
        assert any("python-docx" in str(o) for o in out2)
        import zipfile
        with zipfile.ZipFile(target2) as zf:
            styles = zf.read("word/styles.xml").decode("utf-8")
        assert "微软雅黑" in styles
    finally:
        unwrap()


def test_export_docx_pandoc_missing_fails_explicitly(tmp_path: Path, monkeypatch):
    """2026-08-15: md 输入而 pandoc 不可用(异常环境)时必须显式失败, 不静默
    降级裸 dump——镜像预装 pandoc, 生产不应触发。"""
    handler, root, session, unwrap, mod = _export_docx_harness(tmp_path, monkeypatch)
    try:
        src = root / "notes.md"
        src.parent.mkdir(parents=True, exist_ok=True)
        src.write_text("# 标题\n", encoding="utf-8")
        monkeypatch.setattr(mod, "_run_pandoc", lambda s, t, timeout_s=60: (False, "pandoc not found in this environment"))

        out = list(handler.do_export_docx(
            {"path": "outputs/fail.docx", "source_path": "notes.md"},
            "",
        ))
        target = root / "outputs" / "fail.docx"
        assert not target.exists(), "md without pandoc must not produce a dump docx"
        assert any("pandoc conversion failed" in str(o) for o in out), f"explicit error expected: {out}"
    finally:
        unwrap()


def test_export_docx_txt_with_markdown_content_routes_as_markdown(tmp_path: Path, monkeypatch):
    """审查 A-1: .txt 内容命中 markdown 启发式时必须按 markdown 走 pandoc
    (临时文件 .md 扩展)——保留 .txt 扩展会让 pandoc 按纯文本解析, 标记被
    剥除且无排版(实测: # 标题行 → 无标题样式的纯文本)。"""
    handler, root, session, unwrap, mod = _export_docx_harness(tmp_path, monkeypatch)
    try:
        src = root / "notes.txt"
        src.parent.mkdir(parents=True, exist_ok=True)
        src.write_text("# 标题\n\n- 项一\n- 项二\n", encoding="utf-8")
        calls = []
        fake = _make_fake_pandoc(root / "outputs" / "x.docx")
        monkeypatch.setattr(mod, "_run_pandoc",
                            lambda s, t, timeout_s=60: calls.append((str(s), str(t))) or fake(s, t))

        out = list(handler.do_export_docx(
            {"path": "outputs/txt_md.docx", "source_path": "notes.txt"},
            "",
        ))
        target = root / "outputs" / "txt_md.docx"
        assert target.exists()
        assert len(calls) == 1, f"markdown-looking txt must route to pandoc, calls={calls}"
        tmp_src = Path(calls[0][0])
        assert tmp_src.suffix == ".md", f"markdown-looking txt must use .md ext, got {tmp_src}"
        assert not tmp_src.exists(), "temp source must be cleaned up"
        assert any("pandoc" in str(o) for o in out)
    finally:
        unwrap()


def test_export_docx_gbk_source_decoded(tmp_path: Path, monkeypatch):
    """2026-08-15: GBK/GB2312 编码 txt(Windows 记事本来源)必须正确解码,
    不再依赖模型先 iconv(SOP §4.7 收进工具契约)。"""
    handler, root, session, unwrap, mod = _export_docx_harness(tmp_path, monkeypatch)
    try:
        src = root / "gbk_note.txt"
        src.parent.mkdir(parents=True, exist_ok=True)
        src.write_bytes("工作计划：今天完成报告评审，明天安排代码走查。\n".encode("gb18030"))

        out = list(handler.do_export_docx(
            {"path": "outputs/gbk.docx", "source_path": "gbk_note.txt"},
            "",
        ))
        target = root / "outputs" / "gbk.docx"
        assert target.exists()
        import zipfile
        with zipfile.ZipFile(target) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert "工作计划" in xml, "GBK 中文必须正确解码, 不得乱码"
        assert "完成报告评审" in xml
        assert "�" not in xml, "不得出现替换符"
    finally:
        unwrap()


def test_export_docx_gbk_markdown_source_routed_via_pandoc(tmp_path: Path, monkeypatch):
    """2026-08-15: GBK 编码的 md 走 pandoc 前必须归一为 UTF-8(临时文件),
    且保留 .md 扩展名让 pandoc 正确解析。"""
    handler, root, session, unwrap, mod = _export_docx_harness(tmp_path, monkeypatch)
    try:
        src = root / "gbk_doc.md"
        src.parent.mkdir(parents=True, exist_ok=True)
        src.write_bytes("# 中文标题\n\n正文内容\n".encode("gb18030"))

        calls = []
        fake = _make_fake_pandoc(root / "outputs" / "x.docx")
        monkeypatch.setattr(mod, "_run_pandoc", lambda s, t, timeout_s=60: calls.append((str(s), str(t))) or fake(s, t))

        list(handler.do_export_docx(
            {"path": "outputs/gbk_md.docx", "source_path": "gbk_doc.md"},
            "",
        ))
        target = root / "outputs" / "gbk_md.docx"
        assert target.exists()
        assert len(calls) == 1
        # 临时文件保留 .md 扩展名(否则 pandoc 无法识别格式)且已清理。
        tmp_src = Path(calls[0][0])
        assert tmp_src.suffix == ".md", f"temp source must keep .md ext, got {tmp_src}"
        assert not tmp_src.exists(), "temp source must be cleaned up"
    finally:
        unwrap()
