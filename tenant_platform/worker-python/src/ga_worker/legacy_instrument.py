"""Legacy GenericAgent module instrumentation helpers.

Extracted from managed_agent.py (B3: file size limit). These stateless helpers
monkeypatch legacy modules (agentmain, ga, agent_loop) to enforce RuntimePolicy:
tool schema filtering, dispatch guard, print byte counter, max_turns, handler seed.

Fixes applied during extraction:
- P-M2: dispatch guard and handler seed now return unwrap callables that restore
  originals, preventing leaks across sessions/tasks.
- P-M5 修正(Round16-F5): 模块级 import 保持在外层; 函数体内 import 仅用于
  打破循环依赖(session_files/sop_tool 反向引用本模块)或按需惰性加载,
  不构成重复导入负担。
"""

from __future__ import annotations

import copy
import logging
import re
import sys
from pathlib import Path
from typing import Any, Callable

_LOG = logging.getLogger(__name__)

from ga_worker.limits import ToolPolicy
from ga_worker.session_files import ensure_session_sandbox


def prepare_handler_seed(
    agent: Any,
    seed_working: dict[str, Any],
    agent_factory: Any,
    legacy_mods: dict[str, Any] | None,
) -> Callable[[], None]:
    """Ensure new handler receives restored working; clear carryover handler.

    Returns an unwrap callable that restores the original GenericAgentHandler
    class (P-M2: previously leaked the WrappedHandler across sessions).
    """
    agent.handler = None

    # 项目激活态恢复(审查): working 中保存的 _ga_project_mode_name 在
    # checkpoint 恢复后重新设置到 agent 实例, project_mode hook 才能继续注入。
    project_name = seed_working.get("_ga_project_mode_name") if isinstance(seed_working, dict) else None
    if project_name:
        setattr(agent, "_ga_project_mode_name", project_name)

    if agent_factory is not None and seed_working:
        class _Seed:
            def __init__(self):
                self.working = copy.deepcopy(seed_working)
                self.code_stop_signal: list[int] = []
                self.history_info: list[Any] = []

        agent._adapter_seed_working = copy.deepcopy(seed_working)
        if getattr(agent, "handler", None) is None:
            agent.handler = _Seed()

    original_handler_cls: Any = None
    if legacy_mods and seed_working:
        ga_mod = legacy_mods.get("ga")
        if ga_mod is not None and hasattr(ga_mod, "GenericAgentHandler"):
            original_handler_cls = ga_mod.GenericAgentHandler
            seed = copy.deepcopy(seed_working)

            class WrappedHandler(original_handler_cls):  # type: ignore[misc,valid-type]
                def __init__(self, *args, **kwargs):
                    super().__init__(*args, **kwargs)
                    self.working = copy.deepcopy(seed)

            ga_mod.GenericAgentHandler = WrappedHandler  # type: ignore[assignment]
            agent._adapter_handler_wrapped = original_handler_cls

    def unwrap() -> None:
        if legacy_mods and original_handler_cls is not None:
            ga_mod = legacy_mods.get("ga")
            if ga_mod is not None and getattr(ga_mod, "GenericAgentHandler", None) is not original_handler_cls:
                ga_mod.GenericAgentHandler = original_handler_cls  # type: ignore[assignment]

    return unwrap


def apply_tool_policy(tool_policy: ToolPolicy, legacy_mods: dict[str, Any] | None) -> Any:
    """Filter agentmain.TOOLS_SCHEMA to allowed tools; return previous schema."""
    if legacy_mods is None:
        return None
    agentmain = legacy_mods.get("agentmain")
    if agentmain is None or not hasattr(agentmain, "TOOLS_SCHEMA"):
        return None
    previous = agentmain.TOOLS_SCHEMA
    allowed = tool_policy.allowed_tools
    augmented = list(previous)
    for extra in getattr(agentmain, "_tenant_custom_tools_schema", []) or []:
        if not isinstance(extra, dict):
            continue
        name = extra.get("function", {}).get("name")
        if not name:
            continue
        if any(isinstance(t, dict) and t.get("function", {}).get("name") == name for t in augmented):
            continue
        augmented.append(extra)
    # MCP 工具名由 server_id 动态派生(如 exa__web_search), 静态策略无法
    # 穷举——allowed_tools 中的 "mcp:*" 通配放行全部管理员启用 MCP server
    # 的工具(启用动作本身即管理准入; 保持 deny-by-default)。
    allow_mcp = "mcp:*" in allowed
    mcp_names = frozenset(getattr(agentmain, "_tenant_global_mcp_tool_names", None) or [])
    filtered = [
        t
        for t in augmented
        if isinstance(t, dict)
        and (
            t.get("function", {}).get("name") in allowed
            or (allow_mcp and t.get("function", {}).get("name") in mcp_names)
        )
    ]
    agentmain.TOOLS_SCHEMA = filtered
    return previous


def restore_tool_schema(previous: Any, legacy_mods: dict[str, Any] | None) -> None:
    if previous is None or legacy_mods is None:
        return
    agentmain = legacy_mods.get("agentmain")
    if agentmain is not None:
        agentmain.TOOLS_SCHEMA = previous


_EXPORT_DOCX_TOOL = {
    "type": "function",
    "function": {
        "name": "export_docx",
        "description": "Generate a .docx file inside the session outputs/ sandbox. Markdown/HTML sources (source_path or content) are converted with the built-in pandoc using the built-in Chinese reference template (宋体正文小四/首行缩进2字符/1.5倍行距/黑体标题/西文 Times New Roman/表格浅蓝表头全框线); the optional title becomes the cover Title paragraph (pandoc --metadata=title). The tool verifies the output structure before reporting success. Plain text uses python-docx. For custom typography beyond the preset (e.g. heading numbering, specific heading fonts/sizes), run ../assets/docx_utils.py (make-template with --preset cn or --spec, then md-to-docx) per memory/document_conversion_sop.md — do not hand-write python-docx scripts for standard conversions.",
        "parameters": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "relative output path, usually outputs/<name>.docx"},
                "source_path": {"type": "string", "description": "optional relative md/html/txt source file to convert (pandoc)"},
                "title": {"type": "string", "description": "optional document title; becomes the cover Title paragraph (markdown and plain-text paths)"},
                "content": {"type": "string", "description": "optional text content; markdown content is converted via pandoc; if omitted, the tool reads <file_content> from the assistant reply body"},
            },
            "required": ["path"],
        },
    },
}


# 2026-08-15 模板化升级(社区/官方定案): export_docx 的 pandoc 路径默认套用
# 内置中文模板 assets/reference.docx(构建期 make-template --preset cn 生成:
# 宋体正文小四/首行缩进2字符/1.5倍行距/黑体标题/西文 Times New Roman/表格浅
# 蓝表头全框线)——样式定制走 pandoc 官方 --reference-doc 机制, 不再事后全量
# 强改字体。模板缺失 = 环境异常, 显式报错不降级(与 pandoc 缺失同策略)。
# 纯文本路径无模板(python-docx 默认模板), 仅设 eastAsia=宋体——西文保持
# 默认, 不重蹈全雅黑(西文字形宽重 + 非 Windows 平台回退不可控)覆辙。
_DOCX_TEMPLATE_RELPATH = "assets/reference.docx"
_DOCX_CJK_EAST_ASIA = "宋体"
_MARKDOWN_SOURCE_EXTS = (".md", ".markdown", ".html", ".htm")
_PANDOC_TIMEOUT_S = 60

_MARKDOWN_HINTS = (
    re.compile(r"^#{1,6}\s", re.M),
    re.compile(r"^\s*\|.*\|\s*$", re.M),
    re.compile(r"^\s*[-*]\s+\S", re.M),
    re.compile(r"\*\*[^*]+\*\*"),
    re.compile(r"^```", re.M),
)


def _looks_like_markdown(content: str) -> bool:
    """启发式判定 content 是否为 markdown 结构(标题/表格/列表/粗体/代码块)。"""
    if not content or not content.strip():
        return False
    head = content[:4000]
    return any(p.search(head) for p in _MARKDOWN_HINTS)


def _apply_cjk_east_asia(doc: Any, font_name: str) -> None:
    """仅设 w:eastAsia 中文字体(西文保持模板/默认西文字体——社区惯例:
    中文=宋体/黑体, 西文=Times New Roman/Calibri, 全量雅黑会让西文数字
    字形宽重)。同时清除主题 eastAsia 引用。"""
    from docx.oxml.ns import qn

    for container in (doc.styles.element, doc.element):
        for rPr in container.iter(qn("w:rPr")):
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), font_name)
            if rFonts.get(qn("w:eastAsiaTheme")) is not None:
                del rFonts.attrib[qn("w:eastAsiaTheme")]


def _read_text_robust(path: Path) -> str:
    """编码容错读取(SOP §3.7 收进工具契约): UTF-8 优先, 失败回退 GB18030
    (GBK/GB2312 超集, Windows 记事本来源), 再兜底 errors=replace——不再
    依赖模型记得先 iconv。"""
    data = path.read_bytes()
    for enc in ("utf-8", "gb18030"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _run_pandoc(src: Path, target: Path, timeout_s: int = _PANDOC_TIMEOUT_S,
                reference_doc: Path | None = None,
                metadata_title: str | None = None) -> tuple[bool, str]:
    """镜像内置 pandoc(ga-runner 预装官方二进制)文本类转换。

    reference_doc: 官方 --reference-doc 样式模板(中文排版默认 assets/
    reference.docx, 由调用方解析后传入); metadata_title: --metadata=title:
    (docx Title 样式段落 + docProps, 即"封面标题")——控制字符先清洗(list argv
    无 shell 注入风险, 但 \\n 等会破坏 pandoc 元数据解析)。
    返回 (ok, errmsg)。pandoc 缺失/超时/非零退出/零产物均返回错误——调用方
    决定显式报错(不做静默降级裸 dump)。"""
    import shutil
    import subprocess

    pandoc = shutil.which("pandoc")
    if pandoc is None:
        return False, "pandoc not found in this environment"
    cmd = [pandoc, str(src), "-o", str(target)]
    if reference_doc is not None:
        cmd.append(f"--reference-doc={reference_doc}")
    if metadata_title:
        cleaned = re.sub(r"[\x00-\x1f]", " ", metadata_title).strip()
        if cleaned:
            cmd.append(f"--metadata=title:{cleaned}")
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout_s)
    except subprocess.TimeoutExpired:
        return False, f"pandoc timed out after {timeout_s}s"
    if proc.returncode != 0:
        detail = (proc.stderr or proc.stdout or "").strip()
        if len(detail) > 500:
            detail = detail[:500] + "..."
        return False, f"pandoc failed (exit {proc.returncode}): {detail}"
    if not target.is_file() or target.stat().st_size == 0:
        return False, f"pandoc exited 0 but produced no output: {target}"
    return True, ""


def _docx_template_path(session: Any) -> Path:
    """overlay 内置中文模板绝对路径(LEGACY_ASSETS 物化到 session overlay)。
    缺失=构建期产物缺失, 调用方显式报错。"""
    return Path(session.overlay_dir) / _DOCX_TEMPLATE_RELPATH


def _require_docx_template(session: Any) -> Path:
    """模板缺失显式报错(构建期产物缺失 = 环境异常, 与 pandoc 缺失同策略,
    不静默降级 pandoc 默认模板)。"""
    path = _docx_template_path(session)
    if not path.is_file():
        raise ValueError(
            f"docx reference template missing: {path} "
            "(image build must generate assets/reference.docx)"
        )
    return path


def _verify_docx(target: Path, expect_title: str | None = None) -> tuple[bool, str]:
    """转换后结构化验证(官方/社区共识: 验证闭环必须存在, Anthropic docx skill
    渲染后看图 / SOP §3.8 重读——收进工具契约, 不靠模型自觉): 重读段落/表格,
    title 给定时断言 Title 样式段落存在且含标题文本。失败返回 (False, 原因),
    调用方删除产物并显式报错(避免交付坏文件)。"""
    from docx import Document

    try:
        doc = Document(str(target))
    except Exception as exc:
        return False, f"cannot open docx: {exc}"
    paras = len(doc.paragraphs)
    tables = len(doc.tables)
    # python-docx 的 doc.paragraphs 只含 body 直接子段落(表格内段落不算),
    # 纯表格文档可能 paras==0——以 段落+表格 总和判空(审查: 表格-only 误杀)。
    if paras + tables < 1:
        return False, f"empty document ({paras} paragraphs, {tables} tables)"
    if expect_title:
        # pandoc 把 --metadata=title 当 markdown 解析(smart 引号/强调), 对比前
        # 双侧归一(去标记 + 统一弯引号), 避免合法产物被误删(审查 MAJOR)。
        want = _normalize_title_text(expect_title)
        titles = [_normalize_title_text(p.text)
                  for p in doc.paragraphs if p.style.name.lower() == "title"]
        if not want or not titles or not any(want in t for t in titles):
            return False, f"Title-styled paragraph missing or not containing {expect_title!r}"
    return True, f"{paras} paragraphs, {tables} tables"


def _normalize_title_text(text: str) -> str:
    """标题对比归一: 去除 markdown 标记字符, 统一弯引号/破折号/省略号。
    pandoc --metadata=title 的值按 markdown 解析(smart 扩展默认开), 产物
    Title 段落文本与原始 title 参数存在可预期差异——双侧归一后子串比对。"""
    out = re.sub(r"[#*_`~]", "", text)
    out = out.replace("\u201c", "\"").replace("\u201d", "\"")
    out = out.replace("\u2018", "'").replace("\u2019", "'")
    out = out.replace("\u2026", "...").replace("\u2014", "-")
    return out.strip()

def install_export_docx_tool(session: Any, legacy_mods: dict[str, Any] | None) -> Callable[[], None]:
    """Expose export_docx(round11 审查 I5): GA 原生没有 docx 生成工具, 而
    foundation.session-files.v1 policy 允许 export_docx 且提示语引导 Agent
    调用它——必须提供真实实现(python-docx, 镜像已装), 否则 Agent 会调用
    一个不存在的工具。写入会话 outputs/ 沙箱, 路径逃逸拒绝。"""
    if legacy_mods is None or session is None:
        return lambda: None
    ga_mod = legacy_mods.get("ga")
    agentmain = legacy_mods.get("agentmain")
    if ga_mod is None or agentmain is None:
        return lambda: None
    handler_cls = getattr(ga_mod, "GenericAgentHandler", None)
    if handler_cls is None:
        return lambda: None
    if hasattr(handler_cls, "do_export_docx"):
        raise ValueError("export_docx conflicts with existing handler method")

    from ga_worker.session_files import (
        session_sandbox_root,
        normalize_output_name,
        resolve_under_root,
    )

    runtime_root = Path(session.overlay_dir).parents[1]
    sandbox_root = session_sandbox_root(runtime_root, session.session_key)
    try:
        ensure_session_sandbox(runtime_root, session.session_key)
    except Exception:
        pass

    previous_custom_schema = getattr(agentmain, "_tenant_custom_tools_schema", None)
    custom = list(previous_custom_schema or [])
    for tool in custom:
        if isinstance(tool, dict) and tool.get("function", {}).get("name") == "export_docx":
            raise ValueError("duplicate custom tool name: export_docx")

    def do_export_docx(self: Any, args: dict[str, Any], response: Any) -> Any:
        import re

        step_outcome_cls = getattr(sys.modules.get("agent_loop"), "StepOutcome", None)
        try:
            from docx import Document
        except ImportError:
            yield "[Status] ❌ export_docx 不可用: 缺少 python-docx\n"
            if step_outcome_cls is None:
                raise
            return step_outcome_cls(
                {"status": "error", "message": "python-docx not installed"},
                next_prompt="export_docx unavailable; report the explicit error instead of assuming success.\n",
                should_exit=False,
            )
        try:
            raw_path = str(args.get("path") or "")
            if not raw_path.strip():
                raise ValueError("path is required (e.g. outputs/<name>.docx)")
            normalized = normalize_output_name(raw_path)
            target = resolve_under_root(sandbox_root, normalized)
            target.parent.mkdir(parents=True, exist_ok=True)

            content = str(args.get("content") or "")
            source_path = str(args.get("source_path") or "")
            src_abs: Path | None = None
            if not content and source_path:
                src_abs = resolve_under_root(sandbox_root, source_path)
                if not src_abs.is_file():
                    raise ValueError(f"source_path not found: {source_path}")
            if not content:
                # 与 do_file_write 同构(审查): 优先 response.content, 取最后一个
                # <file_content> 块(模型可能先引用示例再写正文, 首个匹配会取错)。
                body = getattr(response, "content", None)
                if not body:
                    try:
                        body = str(response)
                    except Exception:
                        body = ""
                m = re.findall(r"<file_content[^>]*>(.*?)</file_content>", body or "", re.S)
                if m:
                    content = m[-1].strip()
            if not content.strip() and src_abs is None:
                raise ValueError(
                    "no content to export: pass content, source_path, or <file_content>"
                )

            title = str(args.get("title") or "").strip()
            engine = "python-docx"
            if src_abs is not None:
                ext = src_abs.suffix.lower()
                if ext not in _MARKDOWN_SOURCE_EXTS and ext != ".txt":
                    raise ValueError(
                        f"source_path type not supported for docx: {ext} (use md/html/txt)"
                    )
                # 2026-08-15: 编码容错读入(SOP §3.7 收进工具契约: GBK 不再依赖
                # 模型先 iconv), 统一 UTF-8 落临时文件走 pandoc。审查 A-1:
                # txt 内容命中 markdown 启发式时临时文件必须用 .md 扩展——
                # pandoc 按扩展名推断格式, 保留 .txt 会把 md 当纯文本解析
                # (标记被剥除且无排版, 实测); md/html 保留原扩展名让 pandoc
                # 正确识别格式, 纯文本 txt 走 python-docx。
                content = _read_text_robust(src_abs)
                if ext == ".txt" and not _looks_like_markdown(content):
                    doc = Document()
                    if title:
                        doc.add_heading(title, level=0)
                    for line in content.splitlines():
                        if not line.strip():
                            continue
                        doc.add_paragraph(line.rstrip())
                    _apply_cjk_east_asia(doc, _DOCX_CJK_EAST_ASIA)
                    doc.save(str(target))
                    verified = "plain text"
                else:
                    tmp_ext = ".md" if ext == ".txt" else ext
                    tmp = target.parent / (f".{target.name}.src{tmp_ext}")
                    try:
                        tmp.write_text(content, encoding="utf-8")
                        ok, errmsg = _run_pandoc(
                            tmp, target,
                            reference_doc=_require_docx_template(session),
                            metadata_title=title or None,
                        )
                    finally:
                        try:
                            tmp.unlink(missing_ok=True)
                        except OSError:
                            pass
                    if not ok:
                        raise ValueError(f"pandoc conversion failed: {errmsg}")
                    engine = "pandoc"
                    ok, vmsg = _verify_docx(target, title or None)
                    if not ok:
                        target.unlink(missing_ok=True)
                        raise ValueError(f"docx verification failed: {vmsg}")
                    verified = vmsg
            elif _looks_like_markdown(content):
                # content 是 markdown: 落临时 .md 走 pandoc(保真), 完毕删除。
                tmp = target.parent / (f".{target.name}.src.md")
                try:
                    tmp.write_text(content, encoding="utf-8")
                    ok, errmsg = _run_pandoc(
                        tmp, target,
                        reference_doc=_require_docx_template(session),
                        metadata_title=title or None,
                    )
                finally:
                    try:
                        tmp.unlink(missing_ok=True)
                    except OSError:
                        pass
                if not ok:
                    raise ValueError(f"pandoc conversion failed: {errmsg}")
                engine = "pandoc"
                ok, vmsg = _verify_docx(target, title or None)
                if not ok:
                    target.unlink(missing_ok=True)
                    raise ValueError(f"docx verification failed: {vmsg}")
                verified = vmsg
            else:
                # 纯文本: python-docx 逐行 + 仅设中文 eastAsia 字体(不裸 dump)。
                doc = Document()
                if title:
                    doc.add_heading(title, level=0)
                for line in content.splitlines():
                    if not line.strip():
                        continue
                    doc.add_paragraph(line.rstrip())
                _apply_cjk_east_asia(doc, _DOCX_CJK_EAST_ASIA)
                doc.save(str(target))
                verified = "plain text"
                engine = "python-docx"

            generated = getattr(session, "generated_output_files", None)
            if generated is not None:
                rel = str(target.relative_to(sandbox_root)).replace("\\", "/")
                if rel not in generated:
                    generated.append(rel)
            if engine == "pandoc":
                extra = "模板=assets/reference.docx(内置中文模板)"
                if title:
                    extra += f", 封面标题={title}"
                status_text = (f"[Status] ✅ 已生成 Word 文档(engine=pandoc, {extra}, "
                               f"验证={verified}): {normalized}\n")
            else:
                status_text = (f"[Status] ✅ 已生成 Word 文档(engine=python-docx, "
                               f"验证={verified}): {normalized}\n")
            yield status_text
            # 审查: 非 verbose 模式 yield 的状态文本会被丢弃, 模型收不到任何
            # 反馈(此前 data=None 是盲调用)——返回状态文本作为 tool_result data,
            # 与 do_image_gen/do_file_write 同构; next_prompt 强调 FILE 标记回显。
            return step_outcome_cls(
                status_text,
                next_prompt=(
                    f"已生成 {normalized}。请在最终回复中用 [FILE:{normalized}] "
                    "标记行回显该文件以触发交付(原样保留标记)。\n"
                ),
            )
        except Exception as exc:
            yield f"[Status] ❌ export_docx 失败: {exc}\n"
            step_outcome_cls = getattr(sys.modules.get("agent_loop"), "StepOutcome", None)
            if step_outcome_cls is None:
                raise
            return step_outcome_cls(
                {"status": "error", "message": str(exc)},
                next_prompt="export_docx failed; report the explicit error instead of assuming success.\n",
                should_exit=False,
            )

    original_method = getattr(handler_cls, "do_export_docx", None)
    setattr(handler_cls, "do_export_docx", do_export_docx)
    custom.append(_EXPORT_DOCX_TOOL)
    agentmain._tenant_custom_tools_schema = custom

    def unwrap() -> None:
        if getattr(handler_cls, "do_export_docx", None) is do_export_docx:
            if original_method is None:
                delattr(handler_cls, "do_export_docx")
            else:
                setattr(handler_cls, "do_export_docx", original_method)
        if previous_custom_schema is None:
            if hasattr(agentmain, "_tenant_custom_tools_schema"):
                delattr(agentmain, "_tenant_custom_tools_schema")
        else:
            agentmain._tenant_custom_tools_schema = previous_custom_schema

    return unwrap


_SOPHUB_SEARCH_TOOL = {
    "type": "function",
    "function": {
        "name": "sophub_search",
        "description": "Search approved SOPs in SOPHub via the platform-controlled proxy. Returns matching SOP summaries; use sophub_install to install one into the workspace memory/sops/.",
        "parameters": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "search query"},
            },
            "required": ["query"],
        },
    },
}

_SOPHUB_INSTALL_TOOL = {
    "type": "function",
    "function": {
        "name": "sophub_install",
        "description": "Install an approved SOP from SOPHub into the current workspace memory/sops/<remote_id>.md via the platform-controlled proxy. Never overwrites a locally modified SOP.",
        "parameters": {
            "type": "object",
            "properties": {
                "remote_id": {"type": "string", "description": "SOP remote id from sophub_search results"},
            },
            "required": ["remote_id"],
        },
    },
}


def install_sophub_tools(session: Any, legacy_mods: dict[str, Any] | None) -> Callable[[], None]:
    """Expose sophub_search/sophub_install when the session carries a Sophub
    proxy capability(方案 §5.2 接线, 审查): 工具 schema 追加到
    _tenant_custom_tools_schema, handler 方法安装到 GenericAgentHandler;
    policy 必须允许工具名, dispatch guard 仍会二次拦截。
    """
    if legacy_mods is None or session is None:
        return lambda: None
    ga_mod = legacy_mods.get("ga")
    agentmain = legacy_mods.get("agentmain")
    if ga_mod is None or agentmain is None:
        return lambda: None
    handler_cls = getattr(ga_mod, "GenericAgentHandler", None)
    if handler_cls is None:
        return lambda: None
    proxy = getattr(session, "sophub_proxy", None)
    workspace_memory = getattr(session, "workspace_memory", None)
    if proxy is None or workspace_memory is None:
        return lambda: None

    previous_custom_schema = getattr(agentmain, "_tenant_custom_tools_schema", None)
    original_methods: dict[str, Any] = {}
    installed_methods: dict[str, Any] = {}
    custom = list(previous_custom_schema or [])
    existing_names = {
        tool.get("function", {}).get("name")
        for tool in custom
        if isinstance(tool, dict)
    }
    schemas = [(_SOPHUB_SEARCH_TOOL, "sophub_search"), (_SOPHUB_INSTALL_TOOL, "sophub_install")]
    prepared: list[tuple[str, str]] = []
    for schema, name in schemas:
        if name in existing_names:
            raise ValueError(f"duplicate custom tool name: {name}")
        prepared.append((schema, name))
        existing_names.add(name)

    def make_sophub_handler(proxy: Any, workspace_memory: Path, tool_name: str) -> Callable[..., Any]:
        def do_sophub_tool(self, args, response):
            from ga_worker.sop_tool import SopToolError, sophub_install, sophub_search
            try:
                if tool_name == "sophub_search":
                    result = sophub_search(proxy, args.get("query", ""))
                    yield f"[SOPHub Search]\n{result}\n"
                else:
                    target = sophub_install(proxy, Path(workspace_memory), args.get("remote_id", ""))
                    yield f"[SOPHub Install] installed SOP to {target}\n"
            except SopToolError as exc:
                yield f"[Status] ❌ SOPHub 工具调用失败: {exc.message}\n"
                step_outcome_cls = getattr(sys.modules.get("agent_loop"), "StepOutcome", None)
                if step_outcome_cls is None:
                    raise
                return step_outcome_cls(
                    {"status": "error", "message": str(exc.message)},
                    next_prompt="SOPHub tool failed; report the explicit error instead of assuming success.\n",
                    should_exit=False,
                )
            step_outcome_cls = getattr(sys.modules.get("agent_loop"), "StepOutcome", None)
            if step_outcome_cls is None:
                return
            return step_outcome_cls(None, next_prompt="\n")
        return do_sophub_tool

    for schema, name in prepared:
        method_name = f"do_{name}"
        if hasattr(handler_cls, method_name):
            raise ValueError(f"SOPHub tool conflicts with existing handler method: {name}")
        do_sophub_tool = make_sophub_handler(proxy, workspace_memory, name)
        original_methods[method_name] = getattr(handler_cls, method_name, None)
        installed_methods[method_name] = do_sophub_tool
        setattr(handler_cls, method_name, do_sophub_tool)
        custom.append(schema)

    agentmain._tenant_custom_tools_schema = custom

    def unwrap() -> None:
        for method_name, installed in installed_methods.items():
            if getattr(handler_cls, method_name, None) is not installed:
                continue
            original = original_methods[method_name]
            if original is None:
                delattr(handler_cls, method_name)
            else:
                setattr(handler_cls, method_name, original)
        if previous_custom_schema is None:
            if hasattr(agentmain, "_tenant_custom_tools_schema"):
                delattr(agentmain, "_tenant_custom_tools_schema")
        else:
            agentmain._tenant_custom_tools_schema = previous_custom_schema

    return unwrap




def ensure_session_dirs(session: Any, legacy_mods: dict[str, Any] | None) -> Callable[[], None]:
    """确保会话文件目录(attachments/outputs)存在(Round16-F7 改名:
    原 install_session_file_sandbox 名不副实——它不重定向 cwd 也不装
    sandbox, 只建目录)。

    审查: 不再重定向 handler.cwd 或重写 _get_abs_path——那会破坏 GA 原生
    相对路径语义(./temp、../memory、temp/projects)。容器中 GA handler.cwd
    解析到工作区 temp(进程 workdir 为 GA 根), 附件/输出布局已统一到
    temp/attachments 与 temp/outputs(方案 §6), GA 文件工具直接用原生
    相对路径即可读写 Platform 导入的附件与待交付输出。
    """
    if session is None:
        return lambda: None
    ensure_session_sandbox(Path(session.overlay_dir).parents[1], session.session_key)
    return lambda: None


def install_global_mcp_tools(session: Any, legacy_mods: dict[str, Any] | None) -> Callable[[], None]:
    """Expose the session's administrator-enabled MCP catalog to every tenant task."""
    if legacy_mods is None or session is None:
        return lambda: None
    ga_mod = legacy_mods.get("ga")
    agentmain = legacy_mods.get("agentmain")
    if ga_mod is None or agentmain is None:
        return lambda: None
    handler_cls = getattr(ga_mod, "GenericAgentHandler", None)
    if handler_cls is None:
        return lambda: None

    catalog = getattr(session, "mcp_tools", None) or {}
    if not isinstance(catalog, dict):
        raise ValueError("session mcp_tools must be a mapping")

    previous_custom_schema = getattr(agentmain, "_tenant_custom_tools_schema", None)
    previous_tool_names = getattr(agentmain, "_tenant_global_mcp_tool_names", None)
    original_methods: dict[str, Any] = {}
    installed_methods: dict[str, Any] = {}
    custom = list(previous_custom_schema or [])
    existing_names = {
        tool.get("function", {}).get("name")
        for tool in custom
        if isinstance(tool, dict)
    }
    prepared: list[tuple[str, dict[str, Any], Any, str]] = []

    # Validate the complete catalog before mutating the global legacy handler.
    for ga_name, binding in catalog.items():
        if not isinstance(ga_name, str) or not isinstance(binding, dict):
            raise ValueError("invalid MCP tool binding")
        schema = binding.get("schema")
        client = binding.get("client")
        remote_name = binding.get("tool_name")
        if (
            not isinstance(schema, dict)
            or schema.get("function", {}).get("name") != ga_name
            or not isinstance(remote_name, str)
            or client is None
        ):
            raise ValueError(f"invalid MCP tool binding for {ga_name}")
        method_name = f"do_{ga_name}"
        if hasattr(handler_cls, method_name):
            raise ValueError(f"MCP tool conflicts with existing handler method: {ga_name}")
        if ga_name in existing_names:
            raise ValueError(f"duplicate custom tool name: {ga_name}")
        prepared.append((method_name, schema, client, remote_name))
        existing_names.add(ga_name)

    def make_mcp_handler(client: Any, remote_name: str) -> Callable[..., Any]:
        def do_mcp_tool(self, args, response):
            try:
                public_args = {
                    key: value for key, value in dict(args).items()
                    if key not in ("_index", "_tool_num")
                }
                result = client.call_tool(remote_name, public_args)
            except Exception as exc:
                yield f"[Status] ❌ MCP 工具调用失败: {exc}\n"
                step_outcome_cls = getattr(sys.modules.get("agent_loop"), "StepOutcome", None)
                if step_outcome_cls is None:
                    raise
                return step_outcome_cls(
                    {"status": "error", "message": str(exc)},
                    next_prompt="MCP tool failed; report the explicit error instead of assuming success.\n",
                    should_exit=False,
                )
            yield f"[MCP Result]\n{result}\n"
            step_outcome_cls = getattr(sys.modules.get("agent_loop"), "StepOutcome", None)
            if step_outcome_cls is None:
                return result
            return step_outcome_cls(result, next_prompt="\n")
        return do_mcp_tool

    for method_name, schema, client, remote_name in prepared:
        do_mcp_tool = make_mcp_handler(client, remote_name)
        original_methods[method_name] = getattr(handler_cls, method_name, None)
        installed_methods[method_name] = do_mcp_tool
        setattr(handler_cls, method_name, do_mcp_tool)
        custom.append(schema)

    agentmain._tenant_custom_tools_schema = custom
    agentmain._tenant_global_mcp_tool_names = frozenset(catalog)

    def unwrap() -> None:
        for method_name, installed in installed_methods.items():
            if getattr(handler_cls, method_name, None) is not installed:
                continue
            original = original_methods[method_name]
            if original is None:
                delattr(handler_cls, method_name)
            else:
                setattr(handler_cls, method_name, original)
        if previous_custom_schema is None:
            if hasattr(agentmain, "_tenant_custom_tools_schema"):
                delattr(agentmain, "_tenant_custom_tools_schema")
        else:
            agentmain._tenant_custom_tools_schema = previous_custom_schema
        if previous_tool_names is None:
            if hasattr(agentmain, "_tenant_global_mcp_tool_names"):
                delattr(agentmain, "_tenant_global_mcp_tool_names")
        else:
            agentmain._tenant_global_mcp_tool_names = previous_tool_names

    return unwrap


def install_dispatch_guard(
    tool_policy: ToolPolicy,
    legacy_mods: dict[str, Any] | None,
) -> Callable[[], None]:
    """Install tool-policy dispatch guard on GenericAgentHandler.

    Returns an unwrap callable that restores the original dispatch method
    (P-M2: previously the guard was never restored, leaking across tasks).
    """
    if legacy_mods is None:
        return lambda: None
    ga_mod = legacy_mods.get("ga")
    agent_loop = legacy_mods.get("agent_loop")
    if ga_mod is None:
        return lambda: None
    handler_cls = getattr(ga_mod, "GenericAgentHandler", None)
    if handler_cls is None:
        return lambda: None

    allowed = tool_policy.allowed_tools
    # MCP 工具动态命名, 策略用 mcp:* 通配放行(同 apply_tool_policy)。
    allow_mcp = "mcp:*" in allowed
    agentmain = legacy_mods.get("agentmain")
    mcp_names = frozenset(
        getattr(agentmain, "_tenant_global_mcp_tool_names", None) or []
    ) if agentmain is not None else frozenset()

    # Restore original before (re)installing to avoid double-wrapping.
    original_dispatch = getattr(handler_cls, "_adapter_original_dispatch", None)
    if original_dispatch is None:
        original_dispatch = handler_cls.dispatch
    else:
        handler_cls.dispatch = original_dispatch

    def guarded(self, tool_name, args, response, index=0, tool_num=1):
        if (
            tool_name not in allowed
            and tool_name not in ("no_tool", "bad_json")
            and not (allow_mcp and tool_name in mcp_names)
        ):
            yield f"tool denied by policy: {tool_name}\n"
            step_outcome_cls = getattr(agent_loop, "StepOutcome", None) if agent_loop is not None else None
            if step_outcome_cls is None:
                al_mod = sys.modules.get("agent_loop")
                step_outcome_cls = getattr(al_mod, "StepOutcome", None) if al_mod is not None else None
            if step_outcome_cls is not None:
                return step_outcome_cls(None, next_prompt=f"tool denied: {tool_name}", should_exit=False)
            return None
        return (yield from original_dispatch(self, tool_name, args, response, index=index, tool_num=tool_num))

    handler_cls.dispatch = guarded  # type: ignore[assignment]
    handler_cls._adapter_original_dispatch = original_dispatch  # type: ignore[attr-defined]
    handler_cls._adapter_dispatch_guard = allowed  # type: ignore[attr-defined]

    def unwrap() -> None:
        if getattr(handler_cls, "dispatch", None) is guarded:
            handler_cls.dispatch = original_dispatch  # type: ignore[assignment]
        for attr in ("_adapter_original_dispatch", "_adapter_dispatch_guard"):
            if hasattr(handler_cls, attr):
                delattr(handler_cls, attr)

    return unwrap


def _make_handler_wrapper(
    count_fn: Callable[[str], bool],
    wrapped: list[tuple[Any, Any]],
) -> Callable[[Any], None]:
    """Build a _wrap_handler closure that instruments handler.print.

    Each wrapped handler is recorded in `wrapped` as (handler, original_print)
    so the caller can restore it later.
    """

    def _wrap_handler(handler: Any) -> None:
        if handler is None or getattr(handler, "_adapter_print_wrapped", False):
            return
        original_print = getattr(handler, "print", None)

        def counted_print(*args, **kwargs):
            text = " ".join(str(a) for a in args)
            if count_fn(text):
                return None
            if callable(original_print):
                return original_print(*args, **kwargs)
            return None

        handler.print = counted_print  # type: ignore[assignment]
        handler._adapter_print_wrapped = True  # type: ignore[attr-defined]
        wrapped.append((handler, original_print))

    return _wrap_handler


def _restore_wrapped_handlers(wrapped: list[tuple[Any, Any]]) -> None:
    """Restore original handler.print and clear wrap markers."""
    for handler, original_print in wrapped:
        if not getattr(handler, "_adapter_print_wrapped", False):
            continue
        try:
            if original_print is not None:
                handler.print = original_print  # type: ignore[assignment]
            elif hasattr(handler, "print"):
                delattr(handler, "print")
        except Exception:
            pass
        try:
            delattr(handler, "_adapter_print_wrapped")
        except Exception:
            handler._adapter_print_wrapped = False  # type: ignore[attr-defined]


def _swap_counted_handler(
    ga_mod: Any,
    count_fn: Callable[[str], bool],
    wrap_handler: Callable[[Any], None],
) -> tuple[Any, Any]:
    """Swap ga_mod.GenericAgentHandler to a counted subclass.

    Returns (original_cls, counted_cls). counted_cls is None when no swap
    occurred (already installed for this count_fn, or no target).
    """
    if ga_mod is None or not hasattr(ga_mod, "GenericAgentHandler"):
        return None, None
    original_cls = ga_mod.GenericAgentHandler
    if getattr(original_cls, "_adapter_print_factory", None) is count_fn:
        return original_cls, None

    class _CountedHandler(original_cls):  # type: ignore[misc,valid-type]
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            wrap_handler(self)

    _CountedHandler._adapter_print_factory = count_fn  # type: ignore[attr-defined]
    ga_mod.GenericAgentHandler = _CountedHandler  # type: ignore[assignment]
    return original_cls, _CountedHandler


def _swap_watched_agent(
    agent: Any,
    wrap_handler: Callable[[Any], None],
) -> tuple[Any, Any]:
    """Swap agent.__class__ to a watched subclass that wraps handler on assign.

    Returns (original_class, watched_cls). watched_cls is None when no swap
    occurred (already watched, or __class__ reassignment failed).
    """
    original_class = agent.__class__
    if getattr(agent, "_adapter_handler_watch", False):
        return original_class, None
    agent._adapter_handler_watch = True  # type: ignore[attr-defined]
    try:
        class _WatchedAgent(original_class):  # type: ignore[misc,valid-type]
            def __setattr__(self, name, value):
                super().__setattr__(name, value)
                if name == "handler":
                    wrap_handler(value)

        object.__setattr__(agent, "__class__", _WatchedAgent)
        return original_class, _WatchedAgent
    except Exception as exc:
        # 审查(Minor): 移除隐藏 fallback——此前把 wrap_handler 挂到
        # agent._adapter_wrap_handler 属性, 但没有任何消费方读取它,
        # 类替换失败时输出字节计数会静默失效。显式告警让失效可见。
        _LOG.warning("legacy_instrument: agent class swap failed; "
                     "handler rebuild output counting may be bypassed: %s", exc)
        return original_class, None


def install_handler_print_counter(
    agent: Any,
    count_fn: Callable[[str], bool],
    legacy_mods: dict[str, Any] | None,
) -> Callable[[], None]:
    """Wrap handler.print to count output bytes against the quota.

    Returns an unwrap callable that restores the original agent class, the
    original GenericAgentHandler class, and each wrapped handler's print, so
    counter state does not leak between task runs (P-M2 pattern, matching
    install_dispatch_guard/install_max_turns). 任务即进程(决策 D1)下每任务
    全新 agent, 但 restore 仍必须在同一任务内成对调用, 保证隔离与可重入。
    """
    wrapped: list[tuple[Any, Any]] = []
    wrap_handler = _make_handler_wrapper(count_fn, wrapped)
    wrap_handler(getattr(agent, "handler", None))

    ga_mod = legacy_mods.get("ga") if legacy_mods else None
    original_handler_cls, counted_cls = _swap_counted_handler(ga_mod, count_fn, wrap_handler)
    original_agent_class, watched_cls = _swap_watched_agent(agent, wrap_handler)

    def unwrap() -> None:
        if watched_cls is not None and agent.__class__ is watched_cls:
            object.__setattr__(agent, "__class__", original_agent_class)
        if (counted_cls is not None and ga_mod is not None
                and getattr(ga_mod, "GenericAgentHandler", None) is counted_cls):
            ga_mod.GenericAgentHandler = original_handler_cls  # type: ignore[assignment]
        _restore_wrapped_handlers(wrapped)
        for attr in ("_adapter_handler_watch",):
            if hasattr(agent, attr):
                try:
                    delattr(agent, attr)
                except Exception:
                    pass

    return unwrap


def install_max_turns(
    agent: Any,
    max_turns: int,
    legacy_mods: dict[str, Any] | None,
) -> Callable[[], None]:
    """Force RuntimePolicy.max_turns into agent_runner_loop without editing legacy files.

    Returns an unwrap callable that restores original agent_runner_loop in all
    affected modules (P-M2: previously only unwrapped via finally, now clean).
    """
    unwraps: list[Callable[[], None]] = []

    def _wrap_module(mod: Any, attr: str = "agent_runner_loop") -> None:
        if mod is None or not hasattr(mod, attr):
            return
        original = getattr(mod, attr)
        if getattr(original, "_adapter_max_turns_wrapped", False):
            original._adapter_forced_max_turns = max_turns  # type: ignore[attr-defined]
            return

        def wrapped(*args, **kwargs):
            forced = getattr(wrapped, "_adapter_forced_max_turns", max_turns)
            kwargs = dict(kwargs)
            kwargs["max_turns"] = forced
            return original(*args, **kwargs)

        wrapped._adapter_max_turns_wrapped = True  # type: ignore[attr-defined]
        wrapped._adapter_forced_max_turns = max_turns  # type: ignore[attr-defined]
        setattr(mod, attr, wrapped)

        def unwrap() -> None:
            if getattr(mod, attr, None) is wrapped:
                setattr(mod, attr, original)

        unwraps.append(unwrap)

    mods: list[Any] = []
    if legacy_mods:
        mods.extend([legacy_mods.get("agentmain"), legacy_mods.get("agent_loop")])

    for name in ("agentmain", "agent_loop"):
        mod = sys.modules.get(name)
        if mod is not None and mod not in mods:
            mods.append(mod)
    for mod in mods:
        _wrap_module(mod, "agent_runner_loop")
    agent._adapter_max_turns = max_turns

    def unwrap_all() -> None:
        for u in unwraps:
            try:
                u()
            except Exception:
                pass

    return unwrap_all


def install_image_gen_marker_registry(session: Any, legacy_mods: dict[str, Any] | None) -> Callable[[], None]:
    """Phase B 托管形态(2026-08-14 定稿, 二轮审查 I-2): image_gen 的 marker
    回显兜底——GA do_image_gen 工具返回 [FILE:outputs/...] 标记, 但工具返回
    ≠ 交付发生, 模型必须在最终回复中回显 marker 才触发 Go 侧捕获。本包装
    不替换工具逻辑(ga.py 原生实现), 只在工具成功落盘后把产物相对路径登记
    到 session.generated_output_files——终态时 task_terminal 经
    append_missing_file_markers 自动补写漏回显的 marker(与 export_docx 同款
    机制, legacy_instrument.do_export_docx → task_terminal.py:151-164)。"""
    if legacy_mods is None or session is None:
        return lambda: None
    ga_mod = legacy_mods.get("ga")
    if ga_mod is None:
        return lambda: None
    handler_cls = getattr(ga_mod, "GenericAgentHandler", None)
    if handler_cls is None:
        return lambda: None
    original = getattr(handler_cls, "do_image_gen", None)
    if original is None:
        # 旧版 ga.py 无 image_gen(不兼容): 不安装, 工具不可见属 policy 层职责。
        return lambda: None
    if getattr(original, "_tenant_image_gen_registry", False):
        return lambda: None

    import re

    _MARKER_RE = re.compile(r"\[FILE:(outputs/[^\s\]]+)\]")

    def do_image_gen(self: Any, args: dict[str, Any], response: Any) -> Any:
        gen = original(self, args, response)
        outcome = None
        try:
            while True:
                yield next(gen)
        except StopIteration as stop:
            outcome = stop.value
        if outcome is not None:
            data = getattr(outcome, "data", None)
            text = data if isinstance(data, str) else ""
            if not text:
                text = str(getattr(outcome, "next_prompt", "") or "")
            generated = getattr(session, "generated_output_files", None)
            if generated is not None:
                for marker in _MARKER_RE.findall(text):
                    rel = marker.replace("\\", "/")
                    if rel not in generated:
                        generated.append(rel)
        return outcome

    do_image_gen._tenant_image_gen_registry = True  # type: ignore[attr-defined]
    setattr(handler_cls, "do_image_gen", do_image_gen)

    def unwrap() -> None:
        current = getattr(handler_cls, "do_image_gen", None)
        if current is do_image_gen:
            setattr(handler_cls, "do_image_gen", original)

    return unwrap
